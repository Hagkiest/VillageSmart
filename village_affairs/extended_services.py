import re
from calendar import monthrange
from datetime import datetime, time, timedelta
from decimal import Decimal, InvalidOperation

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from django.db import transaction
from django.db.models import Count, Q, Sum
from django.utils import timezone

from .models import (
    CareObject,
    CareObjectImportBatch,
    DisabledImportBatch,
    DisabledPerson,
    FarmlandImportBatch,
    FarmlandRecord,
    Household,
    LowIncomeImportBatch,
    LowIncomeRecord,
    MediationRecord,
    OrganizationMember,
    PartyFeeRecord,
    PartyMember,
    PartyMemberPositionRecord,
    PartyMemberTransferRecord,
    PublicJobImportBatch,
    PublicJobRecord,
    ProjectImportBatch,
    ProjectRecord,
    ReminderRule,
    Resident,
    RiskCheck,
    RiskCheckImportBatch,
    SubsidyImportBatch,
    SubsidyRecord,
    TodoReminder,
)
from .services import (
    FIELD_LABELS,
    calculate_age,
    generate_error_report_workbook,
    get_import_dir,
    load_excel_rows,
    normalize_identity_number,
    paginate_queryset,
    parse_date_value,
    serialize_scalar,
    validate_excel_filename,
    validate_identity_number,
)


RISK_CHECK_SYSTEM_FIELDS = [
    {'key': 'full_name', 'label': '姓名', 'required': False},
    {'key': 'identity_number', 'label': '身份证号', 'required': True},
    {'key': 'head_name', 'label': '户主姓名', 'required': False},
    {'key': 'head_identity_number', 'label': '户主证件号', 'required': False},
    {'key': 'household_type', 'label': '户属性', 'required': False},
    {'key': 'risk_level', 'label': '风险等级', 'required': False},
    {'key': 'warning_content', 'label': '预警内容', 'required': False},
    {'key': 'medical_amount', 'label': '医疗自付费用金额', 'required': False},
    {'key': 'warning_time', 'label': '预警时间', 'required': False},
]

LOW_INCOME_SYSTEM_FIELDS = [
    {'key': 'full_name', 'label': '居民姓名', 'required': False},
    {'key': 'identity_number', 'label': '身份证号', 'required': True},
    {'key': 'policy_type', 'label': '低收入类型', 'required': False},
    {'key': 'benefit_level', 'label': '享受档次', 'required': False},
    {'key': 'subsidy_amount', 'label': '补贴金额', 'required': False},
    {'key': 'subsidy_cycle', 'label': '补贴周期', 'required': False},
    {'key': 'start_date', 'label': '开始时间', 'required': False},
    {'key': 'end_date', 'label': '结束时间', 'required': False},
    {'key': 'household_member_count', 'label': '全户人数', 'required': False},
    {'key': 'beneficiary_count', 'label': '享受人数', 'required': False},
    {'key': 'household_month_amount', 'label': '户月金额', 'required': False},
    {'key': 'status', 'label': '状态', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]

PROJECT_SYSTEM_FIELDS = [
    {'key': 'project_name', 'label': '项目名称', 'required': True},
    {'key': 'project_source', 'label': '项目库来源', 'required': False},
    {'key': 'project_type', 'label': '项目类型', 'required': False},
    {'key': 'secondary_type', 'label': '二级类型', 'required': False},
    {'key': 'project_status', 'label': '项目状态', 'required': False},
    {'key': 'planning_year', 'label': '规划年度', 'required': False},
    {'key': 'implementation_year', 'label': '实施年度', 'required': False},
    {'key': 'included_in_plan', 'label': '纳入计划', 'required': False},
    {'key': 'total_investment', 'label': '项目预算总投资(万元)', 'required': False},
    {'key': 'settled_amount', 'label': '结算金额(万元)', 'required': False},
    {'key': 'audited_amount', 'label': '决算审计金额(万元)', 'required': False},
    {'key': 'responsible_person', 'label': '督护人/责任人', 'required': False},
    {'key': 'project_location', 'label': '项目地点', 'required': False},
    {'key': 'project_description', 'label': '项目描述', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]

FARMLAND_SYSTEM_FIELDS = [
    {'key': 'plot_code', 'label': '地块编号', 'required': True},
    {'key': 'village_group', 'label': '村组', 'required': False},
    {'key': 'contractor_name', 'label': '承包户', 'required': True},
    {'key': 'contractor_identity_number', 'label': '承包户身份证号', 'required': False},
    {'key': 'linked_resident_id', 'label': '关联居民ID', 'required': False},
    {'key': 'plot_location', 'label': '地块位置', 'required': False},
    {'key': 'area_mu', 'label': '面积（亩）', 'required': False},
    {'key': 'east_boundary', 'label': '东至', 'required': False},
    {'key': 'south_boundary', 'label': '南至', 'required': False},
    {'key': 'west_boundary', 'label': '西至', 'required': False},
    {'key': 'north_boundary', 'label': '北至', 'required': False},
    {'key': 'plot_status', 'label': '地块状态', 'required': False},
    {'key': 'transfer_status', 'label': '流转情况', 'required': False},
    {'key': 'confirmation_status', 'label': '确权情况', 'required': False},
    {'key': 'current_planting', 'label': '当前种植', 'required': False},
    {'key': 'latest_change', 'label': '最新变更', 'required': False},
    {'key': 'change_date', 'label': '变更日期', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]


def suggest_mapping_for_fields(headers, system_fields):
    suggestions = {}
    aliases = {
        'identity_number': ['身份证号', '身份证号码', '证件号码'],
        'head_identity_number': ['户主证件号', '户主身份证号', '户主身份证号码'],
        'full_name': ['姓名', '居民姓名'],
        'policy_type': ['低收入类型', '享受政策类型', '政策类型'],
        'benefit_level': ['享受档次'],
        'subsidy_amount': ['补贴金额'],
        'subsidy_cycle': ['补贴周期'],
        'household_member_count': ['全户人数'],
        'beneficiary_count': ['享受人数'],
        'household_month_amount': ['户月金额'],
        'warning_time': ['预警时间'],
        'disability_type': ['残疾类型'],
        'disability_level': ['残疾等级'],
        'disability_card_number': ['残疾证号', '残疾人证号'],
        'guardian_name': ['监护人姓名'],
        'guardian_phone': ['监护人电话'],
        'issue_date': ['办证日期', '发证日期'],
        'grant_year': ['年度', '年份'],
        'batch_name': ['批次'],
        'subsidy_type': ['补贴类型'],
        'bank_account': ['银行账号', '一卡通账号'],
        'village_group': ['村组', '行政村'],
        'household_population': ['家庭人口', '全户人数'],
        'subsidy_item': ['项目/事项', '项目', '事项'],
        'subsidy_standard': ['规格', '补贴标准'],
        'unit': ['单位'],
        'declared_amount': ['申报金额'],
        'actual_amount': ['实发金额'],
        'payment_status': ['发放状态', '支付状态'],
        'payment_date': ['发放日期', '支付日期'],
        'job_name': ['岗位名称', '岗位'],
        'department': ['主管部门', '用工单位', '管理部门'],
        'start_date': ['开始日期', '合同开始日期', '上岗日期'],
        'end_date': ['结束日期', '合同结束日期', '离岗日期'],
        'required_attendance_days': ['规定出勤天数', '应出勤天数'],
        'actual_attendance_days': ['实际出勤天数'],
        'project_name': ['项目名称'],
        'project_source': ['项目库来源', '项目来源'],
        'project_type': ['项目类型'],
        'secondary_type': ['二级类型', '项目二级类型'],
        'project_status': ['项目状态', '状态'],
        'planning_year': ['规划年度'],
        'implementation_year': ['实施年度'],
        'included_in_plan': ['纳入计划', '是否纳入计划'],
        'total_investment': ['项目预算总投资(万元)', '项目预算总投资', '预算总投资'],
        'settled_amount': ['结算金额(万元)', '结算金额'],
        'audited_amount': ['决算审计金额(万元)', '决算审计金额'],
        'responsible_person': ['督护人/责任人', '责任人', '督护人'],
        'project_location': ['项目地点', '地点'],
        'project_description': ['项目描述', '描述', '项目简介'],
        'plot_code': ['地块编号', '地块编码'],
        'contractor_name': ['承包户', '承包人', '户主姓名'],
        'contractor_identity_number': ['承包户身份证号', '承包户身份证号码', '身份证号', '身份证号码'],
        'linked_resident_id': ['关联居民ID', '居民ID'],
        'plot_location': ['地块位置', '位置'],
        'area_mu': ['面积（亩）', '面积(亩)', '面积', '确权面积'],
        'east_boundary': ['东至'],
        'south_boundary': ['南至'],
        'west_boundary': ['西至'],
        'north_boundary': ['北至'],
        'plot_status': ['地块状态', '状态'],
        'transfer_status': ['流转情况', '是否流转'],
        'confirmation_status': ['确权情况'],
        'current_planting': ['当前种植', '种植作物'],
        'latest_change': ['最新变更', '最新变更情况'],
        'change_date': ['变更日期', '最新变更日期'],
        'care_type': ['关爱类型'],
        'care_level': ['关爱等级'],
        'caregiver_name': ['关爱人员', '帮扶人', '照护人'],
        'caregiver_phone': ['联系方式', '关爱人员电话', '帮扶人电话', '照护人电话'],
        'status': ['状态'],
        'notes': ['备注'],
    }
    for field in system_fields:
        field_key = field['key']
        candidates = [field['label']] + aliases.get(field_key, [])
        for header in headers:
            normalized_header = header.replace(' ', '')
            if any(normalized_header == candidate.replace(' ', '') for candidate in candidates):
                suggestions[field_key] = header
                break
    return suggestions


def normalize_module_mapping(mapping, system_fields):
    allowed_keys = {field['key'] for field in system_fields}
    normalized = {}
    for key, value in (mapping or {}).items():
        if key in allowed_keys and value:
            normalized[key] = str(value).strip()
    return normalized


def map_module_row(row, mapping):
    return {
        field_key: serialize_scalar(row.get(header))
        for field_key, header in mapping.items()
        if header
    }


def create_generic_import_batch(batch_model, uploaded_file):
    validate_excel_filename(uploaded_file.name)
    batch_id = batch_model._meta.pk.default()
    target_path = get_import_dir() / f'{batch_id}_{uploaded_file.name}'
    with target_path.open('wb') as output:
        for chunk in uploaded_file.chunks():
            output.write(chunk)
    headers, rows = load_excel_rows(target_path)
    batch = batch_model.objects.create(
        id=batch_id,
        original_filename=uploaded_file.name,
        file_path=str(target_path),
        source_headers=headers,
        total_rows=len(rows),
    )
    return batch, headers, rows


def parse_decimal_value(value):
    text = serialize_scalar(value)
    if not text:
        return None
    try:
        return Decimal(text)
    except (InvalidOperation, ValueError):
        return None


def parse_int_value(value, default=0):
    text = serialize_scalar(value)
    if not text:
        return default
    try:
        return int(float(text))
    except (TypeError, ValueError):
        return default


def parse_bool_value(value, default=False):
    if value in (None, ''):
        return default
    if isinstance(value, bool):
        return value
    text = serialize_scalar(value).strip().lower()
    if text in {'1', 'true', 'yes', 'y', '是', '纳入', '已纳入'}:
        return True
    if text in {'0', 'false', 'no', 'n', '否', '未纳入'}:
        return False
    return default


def normalize_risk_check_payload(mapped_row):
    identity_number = normalize_identity_number(mapped_row.get('identity_number'))
    resident = Resident.objects.select_related('household').filter(identity_number=identity_number).first() if identity_number else None
    household = resident.household if resident else None
    warning_time = parse_date_value(mapped_row.get('warning_time')) or timezone.localdate()
    medical_amount = parse_decimal_value(mapped_row.get('medical_amount'))
    risk_level = serialize_scalar(mapped_row.get('risk_level'))

    return {
        'resident': resident,
        'full_name': serialize_scalar(mapped_row.get('full_name')) or (resident.full_name if resident else ''),
        'identity_number': identity_number,
        'head_name': serialize_scalar(mapped_row.get('head_name')) or (household.head_name if household else ''),
        'head_identity_number': normalize_identity_number(mapped_row.get('head_identity_number')) or (household.head_identity_number if household else ''),
        'household_type': serialize_scalar(mapped_row.get('household_type')) or (resident.household_type if resident else ''),
        'risk_level': risk_level or '中风险',
        'warning_content': serialize_scalar(mapped_row.get('warning_content')),
        'medical_amount': medical_amount,
        'warning_time': warning_time,
    }


def validate_risk_check_normalized(normalized, original_row):
    errors = []
    id_valid, id_message = validate_identity_number(normalized['identity_number'])
    if not id_valid:
        errors.append(id_message)
    if normalized['risk_level'] and normalized['risk_level'] not in {'高风险', '中风险', '低风险'}:
        errors.append('风险等级必须为高风险、中风险或低风险')
    if serialize_scalar(original_row.get('medical_amount')) and normalized['medical_amount'] is None:
        errors.append('医疗自付费用金额必须是数字')
    return errors


def save_risk_check_from_normalized(normalized):
    defaults = {
        'resident': normalized['resident'],
        'full_name': normalized['full_name'],
        'head_name': normalized['head_name'],
        'head_identity_number': normalized['head_identity_number'],
        'household_type': normalized['household_type'],
        'warning_content': normalized['warning_content'],
        'medical_amount': normalized['medical_amount'],
    }
    item, created = RiskCheck.objects.update_or_create(
        identity_number=normalized['identity_number'],
        risk_level=normalized['risk_level'],
        warning_time=normalized['warning_time'],
        defaults=defaults,
    )
    return item, created


def build_risk_check_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, RISK_CHECK_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')

    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []

    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_risk_check_payload(mapped)
        errors = validate_risk_check_normalized(normalized, mapped)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1

        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'head_name': normalized['head_name'],
                    'household_type': normalized['household_type'],
                    'risk_level': normalized['risk_level'],
                    'warning_content': normalized['warning_content'],
                    'medical_amount': str(normalized['medical_amount']) if normalized['medical_amount'] is not None else '',
                    'warning_time': normalized['warning_time'].isoformat() if normalized['warning_time'] else '',
                    'errors': errors,
                }
            )

    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = RiskCheckImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])

    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_risk_check_import(batch, mapping):
    build_risk_check_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, RISK_CHECK_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []

    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_risk_check_payload(mapped)
        errors = validate_risk_check_normalized(normalized, mapped)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_risk_check_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )

    batch.status = RiskCheckImportBatch.STATUS_IMPORTED if not skipped else RiskCheckImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


def serialize_low_income_record(item, index=None):
    household = item.household
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'full_name': item.full_name,
        'identity_number': item.identity_number,
        'gender': item.gender,
        'ethnicity': item.ethnicity,
        'age': item.resident.age if item.resident else None,
        'phone': item.phone,
        'head_name': item.head_name or (household.head_name if household else ''),
        'relation_to_head': item.relation_to_head,
        'policy_type': item.policy_type,
        'benefit_level': item.benefit_level,
        'subsidy_amount': str(item.subsidy_amount) if item.subsidy_amount is not None else '',
        'subsidy_cycle': item.subsidy_cycle,
        'start_date': item.start_date.isoformat() if item.start_date else '',
        'end_date': item.end_date.isoformat() if item.end_date else '',
        'household_member_count': item.household_member_count,
        'beneficiary_count': item.beneficiary_count,
        'household_month_amount': str(item.household_month_amount) if item.household_month_amount is not None else '',
        'status': item.status,
        'household_no': household.household_no if household else '',
        'village_group': item.village_group or (household.village_group if household else ''),
        'notes': item.notes,
    }


def list_low_income_records(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = LowIncomeRecord.objects.select_related('resident', 'household').all()
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=params['full_name'])
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=params['identity_number'])
    if params.get('policy_type'):
        queryset = queryset.filter(policy_type=params['policy_type'])
    if params.get('status') and params['status'] != '全部':
        queryset = queryset.filter(status=params['status'])

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_low_income_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
    }


def list_low_income_households(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = LowIncomeRecord.objects.select_related('household').all()
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=params['full_name'])
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=params['identity_number'])
    if params.get('policy_type'):
        queryset = queryset.filter(policy_type=params['policy_type'])
    if params.get('status') and params['status'] != '全部':
        queryset = queryset.filter(status=params['status'])

    grouped = {}
    for item in queryset:
        household = item.household
        key = household.id if household else f'resident-{item.identity_number}'
        payload = grouped.setdefault(
            key,
            {
                'id': household.id if household else item.id,
                'village_group': item.village_group or (household.village_group if household else ''),
                'head_name': item.head_name or (household.head_name if household else item.full_name),
                'household_no': household.household_no if household else '',
                'household_member_count': item.household_member_count,
                'beneficiary_count': 0,
                'household_month_amount': Decimal('0'),
                'policy_types': set(),
                'active_count': 0,
                'paused_count': 0,
            },
        )
        payload['beneficiary_count'] += item.beneficiary_count or 0
        payload['household_month_amount'] += item.household_month_amount or Decimal('0')
        if item.policy_type:
            payload['policy_types'].add(item.policy_type)
        if item.status == '在享':
            payload['active_count'] += 1
        else:
            payload['paused_count'] += 1

    ordered = list(grouped.values())
    total = len(ordered)
    start = (page_number - 1) * page_size
    end = start + page_size
    items = []
    for idx, item in enumerate(ordered[start:end], start=start + 1):
        items.append(
            {
                'id': item['id'],
                'seq': idx,
                'village_group': item['village_group'],
                'head_name': item['head_name'],
                'household_no': item['household_no'],
                'household_member_count': item['household_member_count'],
                'beneficiary_count': item['beneficiary_count'],
                'household_month_amount': str(item['household_month_amount']),
                'policy_type': '、'.join(sorted(item['policy_types'])),
                'status_summary': f"在享{item['active_count']}人 / 停享{item['paused_count']}人",
            }
        )
    return {
        'items': items,
        'pagination': {
            'page': page_number,
            'page_size': page_size,
            'total': total,
            'total_pages': max(1, (total + page_size - 1) // page_size),
        },
    }


@transaction.atomic
def create_low_income_record(data):
    resident_id = data.get('resident_id')
    if not resident_id:
        raise ValueError('必须先选择居民')
    resident = Resident.objects.select_related('household').get(id=resident_id)
    household = resident.household
    household_member_count = data.get('household_member_count')
    beneficiary_count = data.get('beneficiary_count')
    subsidy_amount = parse_decimal_value(data.get('subsidy_amount'))
    household_month_amount = parse_decimal_value(data.get('household_month_amount')) or subsidy_amount

    item = LowIncomeRecord.objects.create(
        resident=resident,
        household=household,
        full_name=resident.full_name,
        identity_number=resident.identity_number,
        gender=resident.gender,
        ethnicity=resident.ethnicity,
        phone=resident.phone,
        head_name=household.head_name if household else resident.full_name,
        relation_to_head=resident.relation_to_head,
        village_group=resident.village_group or (household.village_group if household else ''),
        policy_type=serialize_scalar(data.get('policy_type')),
        benefit_level=serialize_scalar(data.get('benefit_level')),
        subsidy_amount=subsidy_amount,
        subsidy_cycle=serialize_scalar(data.get('subsidy_cycle')),
        start_date=parse_date_value(data.get('start_date')),
        end_date=parse_date_value(data.get('end_date')),
        household_member_count=parse_int_value(household_member_count, household.residents.count() if household else 1),
        beneficiary_count=parse_int_value(beneficiary_count, 1),
        household_month_amount=household_month_amount,
        status=serialize_scalar(data.get('status')) or '在享',
        notes=serialize_scalar(data.get('notes')),
    )
    return serialize_low_income_record(item)


def get_low_income_record_detail(record_id):
    item = LowIncomeRecord.objects.select_related('resident', 'household').get(id=record_id)
    return serialize_low_income_record(item)


@transaction.atomic
def update_low_income_record(record_id, data):
    item = LowIncomeRecord.objects.get(id=record_id)
    if 'resident_id' in data and data['resident_id'] != item.resident_id:
        resident = Resident.objects.select_related('household').get(id=data['resident_id'])
        household = resident.household
        item.resident = resident
        item.household = household
        item.full_name = resident.full_name
        item.identity_number = resident.identity_number
        item.gender = resident.gender
        item.ethnicity = resident.ethnicity
        item.phone = resident.phone
        item.head_name = household.head_name if household else resident.full_name
        item.relation_to_head = resident.relation_to_head
        item.village_group = resident.village_group or (household.village_group if household else '')

    if 'policy_type' in data:
        item.policy_type = serialize_scalar(data.get('policy_type'))
    if 'benefit_level' in data:
        item.benefit_level = serialize_scalar(data.get('benefit_level'))
    if 'subsidy_amount' in data:
        item.subsidy_amount = parse_decimal_value(data.get('subsidy_amount'))
    if 'subsidy_cycle' in data:
        item.subsidy_cycle = serialize_scalar(data.get('subsidy_cycle'))
    if 'start_date' in data:
        item.start_date = parse_date_value(data.get('start_date'))
    if 'end_date' in data:
        item.end_date = parse_date_value(data.get('end_date'))
    if 'household_member_count' in data:
        item.household_member_count = parse_int_value(data.get('household_member_count'))
    if 'beneficiary_count' in data:
        item.beneficiary_count = parse_int_value(data.get('beneficiary_count'))
    if 'household_month_amount' in data:
        item.household_month_amount = parse_decimal_value(data.get('household_month_amount'))
    if 'status' in data:
        item.status = serialize_scalar(data.get('status'))
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))

    item.save()
    return serialize_low_income_record(item)


@transaction.atomic
def delete_low_income_record(record_id):
    item = LowIncomeRecord.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def normalize_low_income_payload(mapped_row):
    identity_number = normalize_identity_number(mapped_row.get('identity_number'))
    resident = Resident.objects.select_related('household').filter(identity_number=identity_number).first() if identity_number else None
    household = resident.household if resident else None
    subsidy_amount = parse_decimal_value(mapped_row.get('subsidy_amount'))
    household_month_amount = parse_decimal_value(mapped_row.get('household_month_amount')) or subsidy_amount
    return {
        'resident': resident,
        'household': household,
        'full_name': serialize_scalar(mapped_row.get('full_name')) or (resident.full_name if resident else ''),
        'identity_number': identity_number,
        'policy_type': serialize_scalar(mapped_row.get('policy_type')),
        'benefit_level': serialize_scalar(mapped_row.get('benefit_level')),
        'subsidy_amount': subsidy_amount,
        'subsidy_cycle': serialize_scalar(mapped_row.get('subsidy_cycle')),
        'start_date': parse_date_value(mapped_row.get('start_date')),
        'end_date': parse_date_value(mapped_row.get('end_date')),
        'household_member_count': parse_int_value(mapped_row.get('household_member_count'), household.residents.count() if household else 1),
        'beneficiary_count': parse_int_value(mapped_row.get('beneficiary_count'), 1),
        'household_month_amount': household_month_amount,
        'status': serialize_scalar(mapped_row.get('status')) or '在享',
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_low_income_normalized(normalized, original_row):
    errors = []
    id_valid, id_message = validate_identity_number(normalized['identity_number'])
    if not id_valid:
        errors.append(id_message)
    if normalized['status'] and normalized['status'] not in {'在享', '停享'}:
        errors.append('状态必须为在享或停享')
    if serialize_scalar(original_row.get('subsidy_amount')) and normalized['subsidy_amount'] is None:
        errors.append('补贴金额必须是数字')
    if serialize_scalar(original_row.get('household_month_amount')) and normalized['household_month_amount'] is None:
        errors.append('户月金额必须是数字')
    return errors


def save_low_income_from_normalized(normalized):
    resident = normalized['resident']
    household = normalized['household']
    defaults = {
        'resident': resident,
        'household': household,
        'full_name': normalized['full_name'] or (resident.full_name if resident else ''),
        'gender': resident.gender if resident else '',
        'ethnicity': resident.ethnicity if resident else '',
        'phone': resident.phone if resident else '',
        'head_name': household.head_name if household else (resident.full_name if resident else ''),
        'relation_to_head': resident.relation_to_head if resident else '',
        'village_group': resident.village_group if resident else '',
        'benefit_level': normalized['benefit_level'],
        'subsidy_amount': normalized['subsidy_amount'],
        'subsidy_cycle': normalized['subsidy_cycle'],
        'end_date': normalized['end_date'],
        'household_member_count': normalized['household_member_count'],
        'beneficiary_count': normalized['beneficiary_count'],
        'household_month_amount': normalized['household_month_amount'],
        'status': normalized['status'],
        'notes': normalized['notes'],
    }
    item, created = LowIncomeRecord.objects.update_or_create(
        identity_number=normalized['identity_number'],
        policy_type=normalized['policy_type'],
        start_date=normalized['start_date'],
        defaults=defaults,
    )
    return item, created


def build_low_income_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, LOW_INCOME_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_low_income_payload(mapped)
        errors = validate_low_income_normalized(normalized, mapped)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'policy_type': normalized['policy_type'],
                    'benefit_level': normalized['benefit_level'],
                    'subsidy_amount': str(normalized['subsidy_amount']) if normalized['subsidy_amount'] is not None else '',
                    'subsidy_cycle': normalized['subsidy_cycle'],
                    'start_date': normalized['start_date'].isoformat() if normalized['start_date'] else '',
                    'end_date': normalized['end_date'].isoformat() if normalized['end_date'] else '',
                    'beneficiary_count': normalized['beneficiary_count'],
                    'household_month_amount': str(normalized['household_month_amount']) if normalized['household_month_amount'] is not None else '',
                    'status': normalized['status'],
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = LowIncomeImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_low_income_import(batch, mapping):
    build_low_income_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, LOW_INCOME_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_low_income_payload(mapped)
        errors = validate_low_income_normalized(normalized, mapped)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_low_income_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = LowIncomeImportBatch.STATUS_IMPORTED if not skipped else LowIncomeImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


LOW_INCOME_DETAIL_EXPORT_COLUMNS = [
    ('居民姓名', 'full_name'),
    ('身份证号', 'identity_number'),
    ('性别', 'gender'),
    ('民族', 'ethnicity'),
    ('年龄', 'age'),
    ('联系电话', 'phone'),
    ('户主姓名', 'head_name'),
    ('与户主关系', 'relation_to_head'),
    ('享受政策类型', 'policy_type'),
    ('享受档次', 'benefit_level'),
    ('补贴金额', 'subsidy_amount'),
    ('补贴周期', 'subsidy_cycle'),
    ('开始时间', 'start_date'),
    ('结束时间', 'end_date'),
    ('全户人数', 'household_member_count'),
    ('享受人数', 'beneficiary_count'),
    ('户月金额', 'household_month_amount'),
    ('状态', 'status'),
]

LOW_INCOME_HOUSEHOLD_EXPORT_COLUMNS = [
    ('村组', 'village_group'),
    ('户主姓名', 'head_name'),
    ('户号', 'household_no'),
    ('全户人数', 'household_member_count'),
    ('享受人数', 'beneficiary_count'),
    ('户月金额', 'household_month_amount'),
    ('政策类型', 'policy_type'),
    ('状态统计', 'status_summary'),
]


def build_low_income_export_workbook(view, params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    if view == 'household':
        sheet.title = '低收入按户汇总'
        columns = LOW_INCOME_HOUSEHOLD_EXPORT_COLUMNS
        data = list_low_income_households({**params, 'page': 1, 'page_size': 100000})['items']
    else:
        sheet.title = '低收入人员明细'
        columns = LOW_INCOME_DETAIL_EXPORT_COLUMNS
        data = list_low_income_records({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(columns, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(columns, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def build_error_report_response_batch(batch):
    return generate_error_report_workbook(batch)


DISABLED_SYSTEM_FIELDS = [
    {'key': 'full_name', 'label': '姓名', 'required': False},
    {'key': 'identity_number', 'label': '身份证号', 'required': True},
    {'key': 'gender', 'label': '性别', 'required': False},
    {'key': 'ethnicity', 'label': '民族', 'required': False},
    {'key': 'phone', 'label': '联系电话', 'required': False},
    {'key': 'village_group', 'label': '村组', 'required': False},
    {'key': 'address', 'label': '家庭地址', 'required': False},
    {'key': 'disability_type', 'label': '残疾类型', 'required': False},
    {'key': 'disability_level', 'label': '残疾等级', 'required': False},
    {'key': 'disability_card_number', 'label': '残疾证号', 'required': False},
    {'key': 'issue_date', 'label': '办证日期', 'required': False},
    {'key': 'guardian_name', 'label': '监护人姓名', 'required': False},
    {'key': 'guardian_phone', 'label': '监护人电话', 'required': False},
    {'key': 'status', 'label': '状态', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]


SUBSIDY_SYSTEM_FIELDS = [
    {'key': 'grant_year', 'label': '年度', 'required': True},
    {'key': 'batch_name', 'label': '批次', 'required': False},
    {'key': 'subsidy_type', 'label': '补贴类型', 'required': True},
    {'key': 'full_name', 'label': '姓名', 'required': False},
    {'key': 'identity_number', 'label': '身份证号', 'required': True},
    {'key': 'bank_account', 'label': '银行账号', 'required': False},
    {'key': 'village_group', 'label': '村组', 'required': False},
    {'key': 'household_population', 'label': '家庭人口', 'required': False},
    {'key': 'subsidy_item', 'label': '项目/事项', 'required': False},
    {'key': 'subsidy_standard', 'label': '规格', 'required': False},
    {'key': 'unit', 'label': '单位', 'required': False},
    {'key': 'declared_amount', 'label': '申报金额', 'required': False},
    {'key': 'actual_amount', 'label': '实发金额', 'required': False},
    {'key': 'payment_status', 'label': '发放状态', 'required': False},
    {'key': 'payment_date', 'label': '发放日期', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]

PUBLIC_JOB_SYSTEM_FIELDS = [
    {'key': 'full_name', 'label': '姓名', 'required': False},
    {'key': 'identity_number', 'label': '身份证号', 'required': True},
    {'key': 'job_name', 'label': '岗位名称', 'required': True},
    {'key': 'department', 'label': '主管部门', 'required': False},
    {'key': 'start_date', 'label': '合同开始日期', 'required': False},
    {'key': 'end_date', 'label': '合同结束日期', 'required': False},
    {'key': 'subsidy_amount', 'label': '月补贴标准', 'required': False},
    {'key': 'required_attendance_days', 'label': '规定出勤天数', 'required': False},
    {'key': 'actual_attendance_days', 'label': '实际出勤天数', 'required': False},
    {'key': 'status', 'label': '状态', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]

CARE_OBJECT_SYSTEM_FIELDS = [
    {'key': 'full_name', 'label': '姓名', 'required': False},
    {'key': 'identity_number', 'label': '身份证号', 'required': True},
    {'key': 'gender', 'label': '性别', 'required': False},
    {'key': 'ethnicity', 'label': '民族', 'required': False},
    {'key': 'phone', 'label': '联系电话', 'required': False},
    {'key': 'village_group', 'label': '村组', 'required': False},
    {'key': 'address', 'label': '家庭地址', 'required': False},
    {'key': 'care_type', 'label': '关爱类型', 'required': False},
    {'key': 'care_level', 'label': '关爱等级', 'required': False},
    {'key': 'caregiver_name', 'label': '关爱人员', 'required': False},
    {'key': 'caregiver_phone', 'label': '联系方式', 'required': False},
    {'key': 'notes', 'label': '备注', 'required': False},
]


DISABILITY_TYPE_OPTIONS = {'肢体残疾', '视力残疾', '听力残疾', '言语残疾', '智力残疾', '精神残疾', '多重残疾'}
DISABILITY_LEVEL_OPTIONS = {'一级', '二级', '三级', '四级'}
DISABLED_STATUS_OPTIONS = {'有效', '停用'}
SUBSIDY_TYPE_OPTIONS = {
    '耕地地力保护补贴',
    '水稻补贴',
    '产业奖补',
    '产业发展奖补',
    '危房补助',
    '跨省务工交通补贴',
    '县内务工稳岗补贴',
    '雨露计划补助',
}
SUBSIDY_PAYMENT_STATUS_OPTIONS = {'待发放', '已发放'}
PUBLIC_JOB_STATUS_OPTIONS = {
    PublicJobRecord.STATUS_ON_DUTY,
    PublicJobRecord.STATUS_LEFT,
    PublicJobRecord.STATUS_PENDING,
}


def serialize_disabled_person(item, index=None):
    resident = item.resident
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'full_name': item.full_name,
        'identity_number': item.identity_number,
        'gender': item.gender,
        'ethnicity': item.ethnicity,
        'age': resident.age if resident else None,
        'phone': item.phone,
        'village_group': item.village_group,
        'address': item.address,
        'disability_type': item.disability_type,
        'disability_level': item.disability_level,
        'disability_card_number': item.disability_card_number,
        'issue_date': item.issue_date.isoformat() if item.issue_date else '',
        'guardian_name': item.guardian_name,
        'guardian_phone': item.guardian_phone,
        'status': item.status,
        'notes': item.notes,
    }


def list_disabled_people(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = DisabledPerson.objects.select_related('resident', 'household').all()
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=params['full_name'])
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=params['identity_number'])
    if params.get('disability_type'):
        queryset = queryset.filter(disability_type=params['disability_type'])
    if params.get('disability_level'):
        queryset = queryset.filter(disability_level=params['disability_level'])
    if params.get('status') and params['status'] != '全部':
        queryset = queryset.filter(status=params['status'])

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_disabled_person(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
    }


@transaction.atomic
def create_disabled_person(data):
    resident_id = data.get('resident_id')
    if not resident_id:
        raise ValueError('必须先选择居民')
    resident = Resident.objects.select_related('household').get(id=resident_id)
    household = resident.household
    item, _ = DisabledPerson.objects.update_or_create(
        resident=resident,
        defaults={
            'household': household,
            'full_name': resident.full_name,
            'identity_number': resident.identity_number,
            'gender': resident.gender,
            'ethnicity': resident.ethnicity,
            'phone': resident.phone,
            'village_group': resident.village_group or (household.village_group if household else ''),
            'address': resident.address or (household.address if household else ''),
            'disability_type': serialize_scalar(data.get('disability_type')),
            'disability_level': serialize_scalar(data.get('disability_level')),
            'disability_card_number': serialize_scalar(data.get('disability_card_number')),
            'issue_date': parse_date_value(data.get('issue_date')),
            'guardian_name': serialize_scalar(data.get('guardian_name')),
            'guardian_phone': serialize_scalar(data.get('guardian_phone')),
            'status': serialize_scalar(data.get('status')) or '有效',
            'notes': serialize_scalar(data.get('notes')),
        },
    )
    return serialize_disabled_person(item)


def get_disabled_person_detail(record_id):
    item = DisabledPerson.objects.select_related('resident', 'household').get(id=record_id)
    return serialize_disabled_person(item)


@transaction.atomic
def update_disabled_person(record_id, data):
    item = DisabledPerson.objects.select_related('resident', 'household').get(id=record_id)
    if 'resident_id' in data and data['resident_id'] != item.resident_id:
        resident = Resident.objects.select_related('household').get(id=data['resident_id'])
        household = resident.household
        item.resident = resident
        item.household = household
        item.full_name = resident.full_name
        item.identity_number = resident.identity_number
        item.gender = resident.gender
        item.ethnicity = resident.ethnicity
        item.phone = resident.phone
        item.village_group = resident.village_group or (household.village_group if household else '')
        item.address = resident.address or (household.address if household else '')
    if 'disability_type' in data:
        item.disability_type = serialize_scalar(data.get('disability_type'))
    if 'disability_level' in data:
        item.disability_level = serialize_scalar(data.get('disability_level'))
    if 'disability_card_number' in data:
        item.disability_card_number = serialize_scalar(data.get('disability_card_number'))
    if 'issue_date' in data:
        item.issue_date = parse_date_value(data.get('issue_date'))
    if 'guardian_name' in data:
        item.guardian_name = serialize_scalar(data.get('guardian_name'))
    if 'guardian_phone' in data:
        item.guardian_phone = serialize_scalar(data.get('guardian_phone'))
    if 'status' in data:
        item.status = serialize_scalar(data.get('status')) or '有效'
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    item.save()
    return serialize_disabled_person(item)


@transaction.atomic
def delete_disabled_person(record_id):
    item = DisabledPerson.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def normalize_disabled_payload(mapped_row):
    identity_number = normalize_identity_number(mapped_row.get('identity_number'))
    resident = Resident.objects.select_related('household').filter(identity_number=identity_number).first() if identity_number else None
    household = resident.household if resident else None
    return {
        'resident': resident,
        'household': household,
        'full_name': serialize_scalar(mapped_row.get('full_name')) or (resident.full_name if resident else ''),
        'identity_number': identity_number,
        'gender': serialize_scalar(mapped_row.get('gender')) or (resident.gender if resident else ''),
        'ethnicity': serialize_scalar(mapped_row.get('ethnicity')) or (resident.ethnicity if resident else ''),
        'phone': serialize_scalar(mapped_row.get('phone')) or (resident.phone if resident else ''),
        'village_group': serialize_scalar(mapped_row.get('village_group')) or (resident.village_group if resident else ''),
        'address': serialize_scalar(mapped_row.get('address')) or (resident.address if resident else ''),
        'disability_type': serialize_scalar(mapped_row.get('disability_type')),
        'disability_level': serialize_scalar(mapped_row.get('disability_level')),
        'disability_card_number': serialize_scalar(mapped_row.get('disability_card_number')),
        'issue_date': parse_date_value(mapped_row.get('issue_date')),
        'guardian_name': serialize_scalar(mapped_row.get('guardian_name')),
        'guardian_phone': serialize_scalar(mapped_row.get('guardian_phone')),
        'status': serialize_scalar(mapped_row.get('status')) or '有效',
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_disabled_normalized(normalized):
    errors = []
    id_valid, id_message = validate_identity_number(normalized['identity_number'])
    if not id_valid:
        errors.append(id_message)
    if normalized['disability_type'] and normalized['disability_type'] not in DISABILITY_TYPE_OPTIONS:
        errors.append('残疾类型不在允许范围内')
    if normalized['disability_level'] and normalized['disability_level'] not in DISABILITY_LEVEL_OPTIONS:
        errors.append('残疾等级必须为一级、二级、三级或四级')
    if normalized['status'] and normalized['status'] not in DISABLED_STATUS_OPTIONS:
        errors.append('状态必须为有效或停用')
    return errors


def save_disabled_from_normalized(normalized):
    item, created = DisabledPerson.objects.update_or_create(
        identity_number=normalized['identity_number'],
        defaults={
            'resident': normalized['resident'],
            'household': normalized['household'],
            'full_name': normalized['full_name'],
            'gender': normalized['gender'],
            'ethnicity': normalized['ethnicity'],
            'phone': normalized['phone'],
            'village_group': normalized['village_group'],
            'address': normalized['address'],
            'disability_type': normalized['disability_type'],
            'disability_level': normalized['disability_level'],
            'disability_card_number': normalized['disability_card_number'],
            'issue_date': normalized['issue_date'],
            'guardian_name': normalized['guardian_name'],
            'guardian_phone': normalized['guardian_phone'],
            'status': normalized['status'],
            'notes': normalized['notes'],
        },
    )
    return item, created


def build_disabled_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, DISABLED_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_disabled_payload(mapped)
        errors = validate_disabled_normalized(normalized)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'disability_type': normalized['disability_type'],
                    'disability_level': normalized['disability_level'],
                    'disability_card_number': normalized['disability_card_number'],
                    'guardian_name': normalized['guardian_name'],
                    'guardian_phone': normalized['guardian_phone'],
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = DisabledImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_disabled_import(batch, mapping):
    build_disabled_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, DISABLED_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_disabled_payload(mapped)
        errors = validate_disabled_normalized(normalized)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_disabled_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = DisabledImportBatch.STATUS_IMPORTED if not skipped else DisabledImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


DISABLED_EXPORT_COLUMNS = [
    ('姓名', 'full_name'),
    ('身份证号', 'identity_number'),
    ('性别', 'gender'),
    ('民族', 'ethnicity'),
    ('年龄', 'age'),
    ('联系电话', 'phone'),
    ('村组', 'village_group'),
    ('家庭地址', 'address'),
    ('残疾类型', 'disability_type'),
    ('残疾等级', 'disability_level'),
    ('残疾证号', 'disability_card_number'),
    ('监护人姓名', 'guardian_name'),
    ('监护人电话', 'guardian_phone'),
    ('状态', 'status'),
]


def build_disabled_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '残疾人明细'
    data = list_disabled_people({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(DISABLED_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(DISABLED_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def serialize_organization_member(item, index=None):
    resident = item.resident
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'household_id': item.household_id,
        'org_type': item.org_type,
        'source': item.source,
        'full_name': item.full_name,
        'ethnicity': item.ethnicity,
        'gender': item.gender,
        'identity_number': item.identity_number,
        'birth_date': item.birth_date.isoformat() if item.birth_date else '',
        'age': resident.age if resident else calculate_age(item.birth_date),
        'address': item.address,
        'phone': item.phone,
        'position': item.position,
        'political_status': item.political_status,
        'term_number': item.term_number,
        'term_start': item.term_start.isoformat() if item.term_start else '',
        'term_end': item.term_end.isoformat() if item.term_end else '',
        'status': item.status,
        'notes': item.notes,
    }


def list_organization_members(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = OrganizationMember.objects.select_related('resident', 'household').all()
    if params.get('org_type'):
        queryset = queryset.filter(org_type=params['org_type'])
    if params.get('status'):
        queryset = queryset.filter(status=params['status'])
    if params.get('term_number'):
        queryset = queryset.filter(term_number__icontains=serialize_scalar(params.get('term_number')))
    keyword = serialize_scalar(params.get('keyword'))
    if keyword:
        queryset = queryset.filter(Q(full_name__icontains=keyword) | Q(position__icontains=keyword))
    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_organization_member(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
    }


@transaction.atomic
def create_organization_member(data):
    resident_id = data.get('resident_id')
    if not resident_id:
        raise ValueError('必须先查询并选择居民')
    resident = Resident.objects.select_related('household').get(id=resident_id)
    household = resident.household
    org_type = serialize_scalar(data.get('org_type')) or OrganizationMember.ORG_TYPE_PARTY_BRANCH
    if org_type not in dict(OrganizationMember.ORG_TYPE_CHOICES):
        raise ValueError('组织类型不正确')

    item = OrganizationMember.objects.create(
        resident=resident,
        household=household,
        org_type=org_type,
        source=serialize_scalar(data.get('source')) or '居民档案',
        full_name=resident.full_name,
        ethnicity=resident.ethnicity,
        gender=resident.gender,
        identity_number=resident.identity_number,
        birth_date=resident.birth_date,
        address=resident.address or (household.address if household else ''),
        phone=resident.phone,
        position=serialize_scalar(data.get('position')),
        political_status=serialize_scalar(data.get('political_status')) or resident.political_status,
        term_number=serialize_scalar(data.get('term_number')),
        term_start=parse_date_value(data.get('term_start')),
        term_end=parse_date_value(data.get('term_end')),
        status=serialize_scalar(data.get('status')) or OrganizationMember.STATUS_CURRENT,
        notes=serialize_scalar(data.get('notes')),
    )
    return serialize_organization_member(item)


def get_organization_member_detail(record_id):
    item = OrganizationMember.objects.select_related('resident', 'household').get(id=record_id)
    return serialize_organization_member(item)


@transaction.atomic
def update_organization_member(record_id, data):
    item = OrganizationMember.objects.select_related('resident', 'household').get(id=record_id)
    if 'resident_id' in data and data['resident_id'] != item.resident_id:
        resident = Resident.objects.select_related('household').get(id=data['resident_id'])
        household = resident.household
        item.resident = resident
        item.household = household
        item.full_name = resident.full_name
        item.ethnicity = resident.ethnicity
        item.gender = resident.gender
        item.identity_number = resident.identity_number
        item.birth_date = resident.birth_date
        item.address = resident.address or (household.address if household else '')
        item.phone = resident.phone
        item.political_status = resident.political_status

    if 'org_type' in data:
        org_type = serialize_scalar(data.get('org_type'))
        if org_type and org_type not in dict(OrganizationMember.ORG_TYPE_CHOICES):
            raise ValueError('组织类型不正确')
        item.org_type = org_type or item.org_type
    if 'source' in data:
        item.source = serialize_scalar(data.get('source')) or '居民档案'
    if 'position' in data:
        item.position = serialize_scalar(data.get('position'))
    if 'political_status' in data:
        item.political_status = serialize_scalar(data.get('political_status'))
    if 'term_number' in data:
        item.term_number = serialize_scalar(data.get('term_number'))
    if 'term_start' in data:
        item.term_start = parse_date_value(data.get('term_start'))
    if 'term_end' in data:
        item.term_end = parse_date_value(data.get('term_end'))
    if 'status' in data:
        item.status = serialize_scalar(data.get('status')) or OrganizationMember.STATUS_CURRENT
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    item.save()
    return serialize_organization_member(item)


@transaction.atomic
def delete_organization_member(record_id):
    item = OrganizationMember.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


ORGANIZATION_EXPORT_COLUMNS = [
    ('来源', 'source'),
    ('姓名', 'full_name'),
    ('民族', 'ethnicity'),
    ('性别', 'gender'),
    ('身份证号', 'identity_number'),
    ('出生日期', 'birth_date'),
    ('年龄', 'age'),
    ('家庭住址', 'address'),
    ('联系电话', 'phone'),
    ('职务', 'position'),
    ('政治面貌', 'political_status'),
    ('届数', 'term_number'),
    ('任期开始', 'term_start'),
    ('任期结束', 'term_end'),
    ('状态', 'status'),
    ('备注', 'notes'),
]


def build_organization_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '组织架构成员'
    data = list_organization_members({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(ORGANIZATION_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(ORGANIZATION_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def map_political_status_to_party_member_type(political_status):
    status = serialize_scalar(political_status)
    if not status:
        return None
    if status in {'预备党员'}:
        return PartyMember.MEMBER_TYPE_PROBATIONARY
    if status in {'入党积极分子'}:
        return PartyMember.MEMBER_TYPE_ACTIVE
    if '党员' in status or status == '共产党员':
        return PartyMember.MEMBER_TYPE_FULL
    return None


def sync_party_members_from_residents():
    resident_queryset = Resident.objects.select_related('household').filter(
        Q(political_status__icontains='党员') | Q(political_status__in=['共产党员', '预备党员', '入党积极分子'])
    )
    for resident in resident_queryset:
        member_type = map_political_status_to_party_member_type(resident.political_status)
        if not member_type:
            continue
        household = resident.household
        item, created = PartyMember.objects.get_or_create(
            resident=resident,
            defaults={
                'household': household,
                'source': PartyMember.SOURCE_RESIDENT,
                'full_name': resident.full_name,
                'identity_number': resident.identity_number,
                'gender': resident.gender,
                'birth_date': resident.birth_date,
                'ethnicity': resident.ethnicity,
                'education_level': resident.education_level,
                'phone': resident.phone,
                'address': resident.address or (household.address if household else ''),
                'member_type': member_type,
                'status': PartyMember.STATUS_ACTIVE,
            },
        )
        if created:
            continue
        changed = False
        updates = {
            'household': household,
            'full_name': resident.full_name,
            'identity_number': resident.identity_number,
            'gender': resident.gender,
            'birth_date': resident.birth_date,
            'ethnicity': resident.ethnicity,
            'education_level': resident.education_level,
            'phone': resident.phone,
            'address': resident.address or (household.address if household else ''),
            'member_type': member_type,
        }
        for field, value in updates.items():
            if getattr(item, field) != value:
                setattr(item, field, value)
                changed = True
        if item.source != PartyMember.SOURCE_RESIDENT:
            item.source = PartyMember.SOURCE_RESIDENT
            changed = True
        if changed:
            item.save()


def serialize_party_member_transfer_record(item):
    return {
        'id': item.id,
        'transfer_type': item.transfer_type,
        'transfer_date': item.transfer_date.isoformat() if item.transfer_date else '',
        'from_branch': item.from_branch,
        'to_branch': item.to_branch,
        'reason': item.reason,
        'notes': item.notes,
    }


def serialize_party_member_position_record(item):
    return {
        'id': item.id,
        'branch_name': item.branch_name,
        'position_name': item.position_name,
        'start_date': item.start_date.isoformat() if item.start_date else '',
        'end_date': item.end_date.isoformat() if item.end_date else '',
        'is_current': item.is_current,
        'notes': item.notes,
    }


def format_party_fee_status(item):
    latest_fee = item.fee_records.order_by('-fee_year', '-fee_month', '-id').first()
    if not latest_fee:
        return '未生成'
    return f'{latest_fee.fee_year}-{latest_fee.fee_month:02d} {latest_fee.payment_status}'


def serialize_party_member(item, index=None, include_records=False):
    resident = item.resident
    fee_records_manager = getattr(item, 'fee_records', None)
    age = resident.age if resident else calculate_age(item.birth_date)
    party_age = calculate_age(item.join_party_date)
    payload = {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'household_id': item.household_id,
        'source': item.source,
        'full_name': item.full_name,
        'identity_number': item.identity_number,
        'gender': item.gender,
        'birth_date': item.birth_date.isoformat() if item.birth_date else '',
        'age': age,
        'party_age': party_age,
        'ethnicity': item.ethnicity,
        'education_level': item.education_level,
        'phone': item.phone,
        'address': item.address,
        'member_type': item.member_type,
        'join_party_date': item.join_party_date.isoformat() if item.join_party_date else '',
        'becoming_full_member_date': item.becoming_full_member_date.isoformat() if item.becoming_full_member_date else '',
        'party_branch': item.party_branch,
        'current_position': item.current_position,
        'monthly_party_fee': str(item.monthly_party_fee) if item.monthly_party_fee is not None else '',
        'status': item.status,
        'notes': item.notes,
        'fee_status_label': format_party_fee_status(item) if fee_records_manager is not None else '未生成',
    }
    if include_records:
        payload['transfer_records'] = [serialize_party_member_transfer_record(record) for record in item.transfer_records.all()]
        payload['position_records'] = [serialize_party_member_position_record(record) for record in item.position_records.all()]
    return payload


def list_party_members(params):
    sync_party_members_from_residents()
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = PartyMember.objects.select_related('resident', 'household').prefetch_related('fee_records').all()
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=serialize_scalar(params.get('full_name')))
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=serialize_scalar(params.get('identity_number')))
    if params.get('member_type'):
        queryset = queryset.filter(member_type=serialize_scalar(params.get('member_type')))
    if params.get('party_branch'):
        queryset = queryset.filter(party_branch__icontains=serialize_scalar(params.get('party_branch')))
    if params.get('status') and params.get('status') != '全部':
        queryset = queryset.filter(status=serialize_scalar(params.get('status')))
    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_party_member(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
    }


def build_party_member_base_payload(data, resident=None):
    household = resident.household if resident else None
    source = serialize_scalar(data.get('source')) or (PartyMember.SOURCE_RESIDENT if resident else PartyMember.SOURCE_MANUAL)
    member_type = serialize_scalar(data.get('member_type'))
    if member_type and member_type not in dict(PartyMember.MEMBER_TYPE_CHOICES):
        raise ValueError('人员类别不正确')
    status = serialize_scalar(data.get('status')) or PartyMember.STATUS_ACTIVE
    if status not in dict(PartyMember.STATUS_CHOICES):
        raise ValueError('党员状态不正确')
    identity_number = normalize_identity_number(data.get('identity_number')) or (resident.identity_number if resident else '')
    id_valid, id_message = validate_identity_number(identity_number)
    if not id_valid:
        raise ValueError(id_message)
    full_name = serialize_scalar(data.get('full_name')) or (resident.full_name if resident else '')
    if not full_name:
        raise ValueError('请填写党员姓名')
    return {
        'resident': resident,
        'household': household,
        'source': source,
        'full_name': full_name,
        'identity_number': identity_number,
        'gender': serialize_scalar(data.get('gender')) or (resident.gender if resident else ''),
        'birth_date': parse_date_value(data.get('birth_date')) or (resident.birth_date if resident else None),
        'ethnicity': serialize_scalar(data.get('ethnicity')) or (resident.ethnicity if resident else ''),
        'education_level': serialize_scalar(data.get('education_level')) or (resident.education_level if resident else ''),
        'phone': serialize_scalar(data.get('phone')) or (resident.phone if resident else ''),
        'address': serialize_scalar(data.get('address')) or (resident.address if resident else (household.address if household else '')),
        'member_type': member_type or map_political_status_to_party_member_type(resident.political_status if resident else '') or PartyMember.MEMBER_TYPE_FULL,
        'join_party_date': parse_date_value(data.get('join_party_date')),
        'becoming_full_member_date': parse_date_value(data.get('becoming_full_member_date')),
        'party_branch': serialize_scalar(data.get('party_branch')),
        'current_position': serialize_scalar(data.get('current_position')),
        'monthly_party_fee': parse_decimal_value(data.get('monthly_party_fee')) if 'monthly_party_fee' in data else Decimal('10.00'),
        'status': status,
        'notes': serialize_scalar(data.get('notes')),
    }


def sync_party_member_transfer_records(item, records):
    existing = {record.id: record for record in item.transfer_records.all()}
    keep_ids = []
    for raw_record in records or []:
        record_id = raw_record.get('id')
        transfer_type = serialize_scalar(raw_record.get('transfer_type')) or PartyMemberTransferRecord.TYPE_BRANCH_CHANGE
        if transfer_type not in dict(PartyMemberTransferRecord.TRANSFER_TYPE_CHOICES):
            raise ValueError('流转类型不正确')
        transfer_date = parse_date_value(raw_record.get('transfer_date'))
        if not transfer_date:
            raise ValueError('流转记录必须填写流转日期')
        defaults = {
            'transfer_type': transfer_type,
            'transfer_date': transfer_date,
            'from_branch': serialize_scalar(raw_record.get('from_branch')),
            'to_branch': serialize_scalar(raw_record.get('to_branch')),
            'reason': serialize_scalar(raw_record.get('reason')),
            'notes': serialize_scalar(raw_record.get('notes')),
        }
        if record_id and record_id in existing:
            record = existing[record_id]
            for field, value in defaults.items():
                setattr(record, field, value)
            record.save()
            keep_ids.append(record.id)
        else:
            record = PartyMemberTransferRecord.objects.create(party_member=item, **defaults)
            keep_ids.append(record.id)
    item.transfer_records.exclude(id__in=keep_ids).delete()


def sync_party_member_position_records(item, records):
    existing = {record.id: record for record in item.position_records.all()}
    keep_ids = []
    current_record = None
    for raw_record in records or []:
        position_name = serialize_scalar(raw_record.get('position_name'))
        if not position_name:
            raise ValueError('任职记录必须填写职务名称')
        start_date = parse_date_value(raw_record.get('start_date'))
        end_date = parse_date_value(raw_record.get('end_date'))
        is_current = bool(raw_record.get('is_current'))
        defaults = {
            'branch_name': serialize_scalar(raw_record.get('branch_name')) or item.party_branch,
            'position_name': position_name,
            'start_date': start_date,
            'end_date': end_date,
            'is_current': is_current,
            'notes': serialize_scalar(raw_record.get('notes')),
        }
        record_id = raw_record.get('id')
        if record_id and record_id in existing:
            record = existing[record_id]
            for field, value in defaults.items():
                setattr(record, field, value)
            record.save()
        else:
            record = PartyMemberPositionRecord.objects.create(party_member=item, **defaults)
        keep_ids.append(record.id)
        if is_current and current_record is None:
            current_record = record
    item.position_records.exclude(id__in=keep_ids).delete()
    if current_record:
        item.position_records.exclude(id=current_record.id).update(is_current=False)
        item.current_position = current_record.position_name
        if not item.party_branch and current_record.branch_name:
            item.party_branch = current_record.branch_name
    elif records is not None:
        latest_record = item.position_records.order_by('-is_current', '-start_date', '-id').first()
        item.current_position = latest_record.position_name if latest_record else ''


@transaction.atomic
def create_party_member(data):
    resident = None
    resident_id = data.get('resident_id')
    if resident_id:
        resident = Resident.objects.select_related('household').get(id=resident_id)
    payload = build_party_member_base_payload(data, resident=resident)
    if resident:
        item, created = PartyMember.objects.update_or_create(
            resident=resident,
            defaults=payload,
        )
    else:
        if PartyMember.objects.filter(identity_number=payload['identity_number']).exists():
            raise ValueError('该身份证号已存在党员档案')
        item = PartyMember.objects.create(**payload)
        created = True
    if 'transfer_records' in data:
        sync_party_member_transfer_records(item, data.get('transfer_records') or [])
    if 'position_records' in data:
        sync_party_member_position_records(item, data.get('position_records') or [])
    item.save()
    return serialize_party_member(item, include_records=True), created


def get_party_member_detail(record_id):
    sync_party_members_from_residents()
    item = PartyMember.objects.select_related('resident', 'household').prefetch_related('transfer_records', 'position_records', 'fee_records').get(id=record_id)
    return serialize_party_member(item, include_records=True)


@transaction.atomic
def update_party_member(record_id, data):
    item = PartyMember.objects.select_related('resident', 'household').get(id=record_id)
    resident = item.resident
    if 'resident_id' in data:
        next_resident_id = data.get('resident_id')
        if next_resident_id:
            resident = Resident.objects.select_related('household').get(id=next_resident_id)
            existing = PartyMember.objects.filter(resident=resident).exclude(id=item.id).first()
            if existing:
                raise ValueError('该居民已存在党员档案')
        else:
            resident = None
    payload = build_party_member_base_payload(data, resident=resident)
    for field, value in payload.items():
        setattr(item, field, value)
    item.save()
    if 'transfer_records' in data:
        sync_party_member_transfer_records(item, data.get('transfer_records') or [])
    if 'position_records' in data:
        sync_party_member_position_records(item, data.get('position_records') or [])
    item.save()
    return serialize_party_member(item, include_records=True)


@transaction.atomic
def delete_party_member(record_id):
    item = PartyMember.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


PARTY_MEMBER_EXPORT_COLUMNS = [
    ('姓名', 'full_name'),
    ('身份证号', 'identity_number'),
    ('性别', 'gender'),
    ('出生日期', 'birth_date'),
    ('年龄', 'age'),
    ('党龄', 'party_age'),
    ('民族', 'ethnicity'),
    ('学历', 'education_level'),
    ('入党日期', 'join_party_date'),
    ('转正日期', 'becoming_full_member_date'),
    ('人员类别', 'member_type'),
    ('手机号码', 'phone'),
    ('所在党支部', 'party_branch'),
    ('现居住址', 'address'),
    ('当前职务', 'current_position'),
    ('党费查缴情形', 'fee_status_label'),
    ('状态', 'status'),
]


def build_party_member_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '党员名册'
    normalized_params = {key: params.get(key) for key in params} if hasattr(params, 'get') else dict(params)
    data = list_party_members({**normalized_params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(PARTY_MEMBER_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(PARTY_MEMBER_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def serialize_party_fee_record(item, index=None):
    party_member = item.party_member
    return {
        'id': item.id,
        'seq': index,
        'party_member_id': item.party_member_id,
        'full_name': party_member.full_name,
        'identity_number': party_member.identity_number,
        'party_branch': party_member.party_branch,
        'member_type': party_member.member_type,
        'fee_year': item.fee_year,
        'fee_month': item.fee_month,
        'amount_due': str(item.amount_due),
        'amount_paid': str(item.amount_paid),
        'payment_status': item.payment_status,
        'payment_date': item.payment_date.isoformat() if item.payment_date else '',
        'notes': item.notes,
    }


def list_party_fee_records(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = PartyFeeRecord.objects.select_related('party_member').all()
    if params.get('fee_year'):
        queryset = queryset.filter(fee_year=int(params['fee_year']))
    if params.get('fee_month'):
        queryset = queryset.filter(fee_month=int(params['fee_month']))
    if params.get('party_branch'):
        queryset = queryset.filter(party_member__party_branch__icontains=serialize_scalar(params.get('party_branch')))
    if params.get('full_name'):
        queryset = queryset.filter(party_member__full_name__icontains=serialize_scalar(params.get('full_name')))
    if params.get('payment_status') and params.get('payment_status') != '全部':
        queryset = queryset.filter(payment_status=serialize_scalar(params.get('payment_status')))
    stats_rows = list(queryset.values_list('amount_due', 'amount_paid', 'payment_status'))
    stats = {
        'total_due': str(sum((amount_due or Decimal('0')) for amount_due, _, _ in stats_rows)),
        'total_paid': str(sum((amount_paid or Decimal('0')) for _, amount_paid, _ in stats_rows)),
        'paid_count': sum(1 for _, _, status in stats_rows if status == PartyFeeRecord.PAYMENT_STATUS_PAID),
        'pending_count': sum(1 for _, _, status in stats_rows if status == PartyFeeRecord.PAYMENT_STATUS_PENDING),
    }
    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_party_fee_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'stats': stats,
    }


@transaction.atomic
def generate_party_fee_records(data):
    fee_year = int(data.get('fee_year') or timezone.localdate().year)
    fee_month = int(data.get('fee_month') or timezone.localdate().month)
    if fee_month < 1 or fee_month > 12:
        raise ValueError('党费月份必须在 1 到 12 之间')
    queryset = PartyMember.objects.filter(status=PartyMember.STATUS_ACTIVE).exclude(member_type=PartyMember.MEMBER_TYPE_ACTIVE)
    if data.get('party_branch'):
        queryset = queryset.filter(party_branch__icontains=serialize_scalar(data.get('party_branch')))
    if data.get('full_name'):
        queryset = queryset.filter(full_name__icontains=serialize_scalar(data.get('full_name')))
    created_count = 0
    existing_count = 0
    for member in queryset:
        default_amount = member.monthly_party_fee if member.monthly_party_fee is not None else Decimal('10.00')
        _, created = PartyFeeRecord.objects.get_or_create(
            party_member=member,
            fee_year=fee_year,
            fee_month=fee_month,
            defaults={
                'amount_due': default_amount,
                'amount_paid': Decimal('0.00'),
                'payment_status': PartyFeeRecord.PAYMENT_STATUS_PENDING,
            },
        )
        if created:
            created_count += 1
        else:
            existing_count += 1
    return {
        'message': f'已生成 {fee_year}年{fee_month}月党费记录 {created_count} 条，已存在 {existing_count} 条。',
        'created_count': created_count,
        'existing_count': existing_count,
    }


def get_party_fee_record_detail(record_id):
    item = PartyFeeRecord.objects.select_related('party_member').get(id=record_id)
    return serialize_party_fee_record(item)


@transaction.atomic
def update_party_fee_record(record_id, data):
    item = PartyFeeRecord.objects.select_related('party_member').get(id=record_id)
    if 'amount_due' in data:
        amount_due = parse_decimal_value(data.get('amount_due'))
        item.amount_due = amount_due if amount_due is not None else Decimal('0.00')
    if 'amount_paid' in data:
        amount_paid = parse_decimal_value(data.get('amount_paid'))
        item.amount_paid = amount_paid if amount_paid is not None else Decimal('0.00')
    if 'payment_status' in data:
        payment_status = serialize_scalar(data.get('payment_status')) or PartyFeeRecord.PAYMENT_STATUS_PENDING
        if payment_status not in dict(PartyFeeRecord.PAYMENT_STATUS_CHOICES):
            raise ValueError('党费缴纳状态不正确')
        item.payment_status = payment_status
        if payment_status == PartyFeeRecord.PAYMENT_STATUS_PAID and item.amount_paid == Decimal('0.00'):
            item.amount_paid = item.amount_due
            if not item.payment_date:
                item.payment_date = timezone.localdate()
    if 'payment_date' in data:
        item.payment_date = parse_date_value(data.get('payment_date'))
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    item.save()
    return serialize_party_fee_record(item)


@transaction.atomic
def mark_party_fee_paid(record_id):
    item = PartyFeeRecord.objects.get(id=record_id)
    item.payment_status = PartyFeeRecord.PAYMENT_STATUS_PAID
    item.amount_paid = item.amount_due
    item.payment_date = item.payment_date or timezone.localdate()
    item.save(update_fields=['payment_status', 'amount_paid', 'payment_date', 'updated_at'])
    return serialize_party_fee_record(item)


PARTY_FEE_EXPORT_COLUMNS = [
    ('姓名', 'full_name'),
    ('身份证号', 'identity_number'),
    ('所在党支部', 'party_branch'),
    ('人员类别', 'member_type'),
    ('年度', 'fee_year'),
    ('月份', 'fee_month'),
    ('应缴金额', 'amount_due'),
    ('实缴金额', 'amount_paid'),
    ('缴纳状态', 'payment_status'),
    ('缴纳日期', 'payment_date'),
    ('备注', 'notes'),
]


def build_party_fee_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '党费缴纳'
    normalized_params = {key: params.get(key) for key in params} if hasattr(params, 'get') else dict(params)
    data = list_party_fee_records({**normalized_params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(PARTY_FEE_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(PARTY_FEE_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def serialize_subsidy_record(item, index=None):
    resident = item.resident
    household = item.household or (resident.household if resident and resident.household else None)
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'grant_year': item.grant_year,
        'batch_name': item.batch_name,
        'subsidy_type': item.subsidy_type,
        'full_name': item.full_name,
        'identity_number': item.identity_number,
        'gender': resident.gender if resident else '',
        'age': resident.age if resident else None,
        'phone': resident.phone if resident else '',
        'head_name': household.head_name if household else '',
        'relation_to_head': resident.relation_to_head if resident else '',
        'household_no': household.household_no if household else '',
        'bank_account': item.bank_account,
        'village_group': item.village_group,
        'household_population': item.household_population,
        'subsidy_item': item.subsidy_item,
        'subsidy_standard': item.subsidy_standard,
        'unit': item.unit,
        'declared_amount': str(item.declared_amount) if item.declared_amount is not None else '',
        'actual_amount': str(item.actual_amount) if item.actual_amount is not None else '',
        'payment_status': item.payment_status,
        'payment_date': item.payment_date.isoformat() if item.payment_date else '',
        'notes': item.notes,
    }


def list_subsidy_records(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = SubsidyRecord.objects.select_related('resident', 'household').all()
    if params.get('grant_year'):
        queryset = queryset.filter(grant_year=int(params['grant_year']))
    if params.get('subsidy_type'):
        queryset = queryset.filter(subsidy_type=params['subsidy_type'])
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=params['full_name'])
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=params['identity_number'])
    if params.get('batch_name'):
        queryset = queryset.filter(batch_name__icontains=params['batch_name'])
    if params.get('village_group'):
        queryset = queryset.filter(village_group__icontains=params['village_group'])
    if params.get('payment_status'):
        queryset = queryset.filter(payment_status=params['payment_status'])
    if params.get('keyword'):
        keyword = serialize_scalar(params.get('keyword'))
        queryset = queryset.filter(
            Q(full_name__icontains=keyword)
            | Q(identity_number__icontains=keyword)
            | Q(bank_account__icontains=keyword)
            | Q(subsidy_item__icontains=keyword)
        )
    if params.get('resident_id'):
        queryset = queryset.filter(resident_id=params['resident_id'])

    stats_queryset = list(queryset.values_list('identity_number', 'declared_amount', 'actual_amount'))
    beneficiary_count = len({identity for identity, _, _ in stats_queryset if identity})
    record_count = len(stats_queryset)
    declared_amount = sum((amount or Decimal('0')) for _, amount, _ in stats_queryset)
    actual_amount = sum((amount or Decimal('0')) for _, _, amount in stats_queryset)

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_subsidy_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'stats': {
            'beneficiary_count': beneficiary_count,
            'record_count': record_count,
            'declared_amount': str(declared_amount),
            'actual_amount': str(actual_amount),
        },
    }


@transaction.atomic
def create_subsidy_record(data):
    resident_id = data.get('resident_id')
    if not resident_id:
        raise ValueError('必须先选择居民')
    resident = Resident.objects.select_related('household').get(id=resident_id)
    household = resident.household
    disabled_person = DisabledPerson.objects.filter(resident=resident).first()
    grant_year = parse_int_value(data.get('grant_year'), timezone.localdate().year)
    declared_amount = parse_decimal_value(data.get('declared_amount'))
    actual_amount = parse_decimal_value(data.get('actual_amount')) or declared_amount
    item = SubsidyRecord.objects.create(
        resident=resident,
        disabled_person=disabled_person,
        household=household,
        grant_year=grant_year,
        batch_name=serialize_scalar(data.get('batch_name')),
        subsidy_type=serialize_scalar(data.get('subsidy_type')),
        full_name=resident.full_name,
        identity_number=resident.identity_number,
        bank_account=serialize_scalar(data.get('bank_account')) or resident.bank_account,
        village_group=serialize_scalar(data.get('village_group')) or resident.village_group,
        household_population=parse_int_value(data.get('household_population'), household.residents.count() if household else 1),
        subsidy_item=serialize_scalar(data.get('subsidy_item')) or serialize_scalar(data.get('subsidy_type')),
        subsidy_standard=serialize_scalar(data.get('subsidy_standard')),
        unit=serialize_scalar(data.get('unit')),
        declared_amount=declared_amount,
        actual_amount=actual_amount,
        payment_status=serialize_scalar(data.get('payment_status')) or '待发放',
        payment_date=parse_date_value(data.get('payment_date')),
        notes=serialize_scalar(data.get('notes')),
    )
    return serialize_subsidy_record(item)


def get_subsidy_record_detail(record_id):
    item = SubsidyRecord.objects.select_related('resident', 'household').get(id=record_id)
    return serialize_subsidy_record(item)


@transaction.atomic
def update_subsidy_record(record_id, data):
    item = SubsidyRecord.objects.select_related('resident', 'household').get(id=record_id)
    if 'resident_id' in data and data['resident_id'] != item.resident_id:
        resident = Resident.objects.select_related('household').get(id=data['resident_id'])
        household = resident.household
        item.resident = resident
        item.disabled_person = DisabledPerson.objects.filter(resident=resident).first()
        item.household = household
        item.full_name = resident.full_name
        item.identity_number = resident.identity_number
        item.bank_account = resident.bank_account
        item.village_group = resident.village_group
        item.household_population = household.residents.count() if household else 1
    if 'grant_year' in data:
        item.grant_year = parse_int_value(data.get('grant_year'), timezone.localdate().year)
    if 'batch_name' in data:
        item.batch_name = serialize_scalar(data.get('batch_name'))
    if 'subsidy_type' in data:
        item.subsidy_type = serialize_scalar(data.get('subsidy_type'))
    if 'bank_account' in data:
        item.bank_account = serialize_scalar(data.get('bank_account'))
    if 'village_group' in data:
        item.village_group = serialize_scalar(data.get('village_group'))
    if 'household_population' in data:
        item.household_population = parse_int_value(data.get('household_population'), 1)
    if 'subsidy_item' in data:
        item.subsidy_item = serialize_scalar(data.get('subsidy_item'))
    if 'subsidy_standard' in data:
        item.subsidy_standard = serialize_scalar(data.get('subsidy_standard'))
    if 'unit' in data:
        item.unit = serialize_scalar(data.get('unit'))
    if 'declared_amount' in data:
        item.declared_amount = parse_decimal_value(data.get('declared_amount'))
    if 'actual_amount' in data:
        item.actual_amount = parse_decimal_value(data.get('actual_amount'))
    if 'payment_status' in data:
        item.payment_status = serialize_scalar(data.get('payment_status')) or '待发放'
    if 'payment_date' in data:
        item.payment_date = parse_date_value(data.get('payment_date'))
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    item.save()
    return serialize_subsidy_record(item)


@transaction.atomic
def delete_subsidy_record(record_id):
    item = SubsidyRecord.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def normalize_subsidy_payload(mapped_row):
    identity_number = normalize_identity_number(mapped_row.get('identity_number'))
    resident = Resident.objects.select_related('household').filter(identity_number=identity_number).first() if identity_number else None
    household = resident.household if resident else None
    disabled_person = DisabledPerson.objects.filter(resident=resident).first() if resident else None
    declared_amount = parse_decimal_value(mapped_row.get('declared_amount'))
    actual_amount = parse_decimal_value(mapped_row.get('actual_amount')) or declared_amount
    return {
        'resident': resident,
        'disabled_person': disabled_person,
        'household': household,
        'grant_year': parse_int_value(mapped_row.get('grant_year'), timezone.localdate().year),
        'batch_name': serialize_scalar(mapped_row.get('batch_name')),
        'subsidy_type': serialize_scalar(mapped_row.get('subsidy_type')),
        'full_name': serialize_scalar(mapped_row.get('full_name')) or (resident.full_name if resident else ''),
        'identity_number': identity_number,
        'bank_account': serialize_scalar(mapped_row.get('bank_account')) or (resident.bank_account if resident else ''),
        'village_group': serialize_scalar(mapped_row.get('village_group')) or (resident.village_group if resident else ''),
        'household_population': parse_int_value(mapped_row.get('household_population'), household.residents.count() if household else 1),
        'subsidy_item': serialize_scalar(mapped_row.get('subsidy_item')) or serialize_scalar(mapped_row.get('subsidy_type')),
        'subsidy_standard': serialize_scalar(mapped_row.get('subsidy_standard')),
        'unit': serialize_scalar(mapped_row.get('unit')),
        'declared_amount': declared_amount,
        'actual_amount': actual_amount,
        'payment_status': serialize_scalar(mapped_row.get('payment_status')) or '待发放',
        'payment_date': parse_date_value(mapped_row.get('payment_date')),
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_subsidy_normalized(normalized, original_row):
    errors = []
    id_valid, id_message = validate_identity_number(normalized['identity_number'])
    if not id_valid:
        errors.append(id_message)
    if normalized['grant_year'] <= 0:
        errors.append('年度必须为正整数')
    if not normalized['full_name']:
        errors.append('姓名不能为空')
    if normalized['subsidy_type'] not in SUBSIDY_TYPE_OPTIONS:
        errors.append('补贴类型不在系统支持范围内')
    if normalized['payment_status'] and normalized['payment_status'] not in SUBSIDY_PAYMENT_STATUS_OPTIONS:
        errors.append('发放状态必须为待发放或已发放')
    if serialize_scalar(original_row.get('declared_amount')) and normalized['declared_amount'] is None:
        errors.append('申报金额必须是数字')
    if serialize_scalar(original_row.get('actual_amount')) and normalized['actual_amount'] is None:
        errors.append('实发金额必须是数字')
    return errors


def save_subsidy_from_normalized(normalized):
    item, created = SubsidyRecord.objects.update_or_create(
        identity_number=normalized['identity_number'],
        grant_year=normalized['grant_year'],
        subsidy_type=normalized['subsidy_type'],
        batch_name=normalized['batch_name'],
        defaults={
            'resident': normalized['resident'],
            'disabled_person': normalized['disabled_person'],
            'household': normalized['household'],
            'full_name': normalized['full_name'],
            'bank_account': normalized['bank_account'],
            'village_group': normalized['village_group'],
            'household_population': normalized['household_population'],
            'subsidy_item': normalized['subsidy_item'],
            'subsidy_standard': normalized['subsidy_standard'],
            'unit': normalized['unit'],
            'declared_amount': normalized['declared_amount'],
            'actual_amount': normalized['actual_amount'],
            'payment_status': normalized['payment_status'],
            'payment_date': normalized['payment_date'],
            'notes': normalized['notes'],
        },
    )
    return item, created


def build_subsidy_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, SUBSIDY_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_subsidy_payload(mapped)
        errors = validate_subsidy_normalized(normalized, mapped)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'grant_year': normalized['grant_year'],
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'subsidy_type': normalized['subsidy_type'],
                    'declared_amount': str(normalized['declared_amount']) if normalized['declared_amount'] is not None else '',
                    'actual_amount': str(normalized['actual_amount']) if normalized['actual_amount'] is not None else '',
                    'payment_status': normalized['payment_status'],
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = SubsidyImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_subsidy_import(batch, mapping):
    build_subsidy_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, SUBSIDY_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_subsidy_payload(mapped)
        errors = validate_subsidy_normalized(normalized, mapped)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_subsidy_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = SubsidyImportBatch.STATUS_IMPORTED if not skipped else SubsidyImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


SUBSIDY_EXPORT_COLUMNS = [
    ('年度', 'grant_year'),
    ('批次', 'batch_name'),
    ('补贴类型', 'subsidy_type'),
    ('姓名', 'full_name'),
    ('身份证号', 'identity_number'),
    ('银行账号', 'bank_account'),
    ('村组', 'village_group'),
    ('家庭人口', 'household_population'),
    ('项目/事项', 'subsidy_item'),
    ('规格', 'subsidy_standard'),
    ('单位', 'unit'),
    ('申报金额', 'declared_amount'),
    ('实发金额', 'actual_amount'),
    ('发放状态', 'payment_status'),
    ('发放日期', 'payment_date'),
]


def build_subsidy_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '政策性补贴台账'
    data = list_subsidy_records({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(SUBSIDY_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(SUBSIDY_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def get_latest_low_income_record(resident):
    if not resident:
        return None
    return LowIncomeRecord.objects.filter(resident=resident).order_by('-updated_at', '-id').first()


def format_contract_period(start_date, end_date):
    if start_date and end_date:
        return f'{start_date.isoformat()} 至 {end_date.isoformat()}'
    if start_date:
        return f'{start_date.isoformat()} 起'
    if end_date:
        return f'截至 {end_date.isoformat()}'
    return ''


def serialize_public_job_record(item, index=None):
    resident = item.resident
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'full_name': item.full_name,
        'identity_number': item.identity_number,
        'gender': item.gender,
        'age': resident.age if resident else None,
        'phone': item.phone,
        'village_group': item.village_group,
        'head_name': item.head_name,
        'household_type': item.household_type,
        'low_income_type': item.low_income_type,
        'job_name': item.job_name,
        'department': item.department,
        'start_date': item.start_date.isoformat() if item.start_date else '',
        'end_date': item.end_date.isoformat() if item.end_date else '',
        'contract_period': format_contract_period(item.start_date, item.end_date),
        'subsidy_amount': str(item.subsidy_amount) if item.subsidy_amount is not None else '',
        'required_attendance_days': item.required_attendance_days,
        'actual_attendance_days': item.actual_attendance_days,
        'status': item.status,
        'notes': item.notes,
    }


def list_public_job_records(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = PublicJobRecord.objects.select_related('resident', 'household', 'low_income_record').all()
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=params['full_name'])
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=params['identity_number'])
    if params.get('job_name'):
        queryset = queryset.filter(job_name=params['job_name'])
    if params.get('department'):
        queryset = queryset.filter(department=params['department'])
    if params.get('status'):
        queryset = queryset.filter(status=params['status'])
    if params.get('resident_id'):
        queryset = queryset.filter(resident_id=params['resident_id'])

    stats = queryset.aggregate(subsidy_total=Sum('subsidy_amount'))
    active_count = queryset.filter(status=PublicJobRecord.STATUS_ON_DUTY).count()
    record_count = queryset.count()
    job_type_count = queryset.exclude(job_name='').values('job_name').distinct().count()

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_public_job_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'stats': {
            'active_count': active_count,
            'record_count': record_count,
            'job_type_count': job_type_count,
            'subsidy_total': str(stats['subsidy_total'] or Decimal('0.00')),
        },
        'filter_options': {
            'job_names': list(
                PublicJobRecord.objects.exclude(job_name='').values_list('job_name', flat=True).distinct().order_by('job_name')
            ),
            'departments': list(
                PublicJobRecord.objects.exclude(department='').values_list('department', flat=True).distinct().order_by('department')
            ),
        },
    }


@transaction.atomic
def create_public_job_record(data):
    resident_id = data.get('resident_id')
    if not resident_id:
        raise ValueError('必须先选择居民')
    job_name = serialize_scalar(data.get('job_name'))
    status = serialize_scalar(data.get('status')) or PublicJobRecord.STATUS_ON_DUTY
    if not job_name:
        raise ValueError('请填写岗位名称')
    if status not in PUBLIC_JOB_STATUS_OPTIONS:
        raise ValueError('状态必须为在岗、离岗或待岗')
    resident = Resident.objects.select_related('household').get(id=resident_id)
    household = resident.household
    low_income_record = get_latest_low_income_record(resident)
    start_date = parse_date_value(data.get('start_date'))
    end_date = parse_date_value(data.get('end_date'))
    if start_date and end_date and end_date < start_date:
        raise ValueError('合同结束日期不能早于开始日期')
    item = PublicJobRecord.objects.create(
        resident=resident,
        household=household,
        low_income_record=low_income_record,
        full_name=resident.full_name,
        identity_number=resident.identity_number,
        gender=resident.gender,
        phone=resident.phone,
        head_name=(household.head_name if household else ''),
        village_group=serialize_scalar(data.get('village_group')) or resident.village_group or (household.village_group if household else ''),
        household_type=resident.household_type or (household.household_type if household else ''),
        low_income_type=(low_income_record.policy_type if low_income_record else ''),
        job_name=job_name,
        department=serialize_scalar(data.get('department')),
        start_date=start_date,
        end_date=end_date,
        subsidy_amount=parse_decimal_value(data.get('subsidy_amount')),
        required_attendance_days=parse_int_value(data.get('required_attendance_days')),
        actual_attendance_days=parse_int_value(data.get('actual_attendance_days')),
        status=status,
        notes=serialize_scalar(data.get('notes')),
    )
    return serialize_public_job_record(item)


def get_public_job_record_detail(record_id):
    item = PublicJobRecord.objects.select_related('resident', 'household', 'low_income_record').get(id=record_id)
    return serialize_public_job_record(item)


@transaction.atomic
def update_public_job_record(record_id, data):
    item = PublicJobRecord.objects.select_related('resident', 'household', 'low_income_record').get(id=record_id)
    if 'resident_id' in data and data['resident_id'] != item.resident_id:
        resident = Resident.objects.select_related('household').get(id=data['resident_id'])
        household = resident.household
        low_income_record = get_latest_low_income_record(resident)
        item.resident = resident
        item.household = household
        item.low_income_record = low_income_record
        item.full_name = resident.full_name
        item.identity_number = resident.identity_number
        item.gender = resident.gender
        item.phone = resident.phone
        item.head_name = household.head_name if household else ''
        item.village_group = resident.village_group or (household.village_group if household else '')
        item.household_type = resident.household_type or (household.household_type if household else '')
        item.low_income_type = low_income_record.policy_type if low_income_record else ''
    if 'job_name' in data:
        item.job_name = serialize_scalar(data.get('job_name'))
        if not item.job_name:
            raise ValueError('请填写岗位名称')
    if 'department' in data:
        item.department = serialize_scalar(data.get('department'))
    if 'start_date' in data:
        item.start_date = parse_date_value(data.get('start_date'))
    if 'end_date' in data:
        item.end_date = parse_date_value(data.get('end_date'))
    if 'subsidy_amount' in data:
        item.subsidy_amount = parse_decimal_value(data.get('subsidy_amount'))
    if 'required_attendance_days' in data:
        item.required_attendance_days = parse_int_value(data.get('required_attendance_days'))
    if 'actual_attendance_days' in data:
        item.actual_attendance_days = parse_int_value(data.get('actual_attendance_days'))
    if 'status' in data:
        item.status = serialize_scalar(data.get('status')) or PublicJobRecord.STATUS_ON_DUTY
        if item.status not in PUBLIC_JOB_STATUS_OPTIONS:
            raise ValueError('状态必须为在岗、离岗或待岗')
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    if item.start_date and item.end_date and item.end_date < item.start_date:
        raise ValueError('合同结束日期不能早于开始日期')
    item.save()
    return serialize_public_job_record(item)


@transaction.atomic
def delete_public_job_record(record_id):
    item = PublicJobRecord.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def normalize_public_job_payload(mapped_row):
    identity_number = normalize_identity_number(mapped_row.get('identity_number'))
    resident = Resident.objects.select_related('household').filter(identity_number=identity_number).first() if identity_number else None
    household = resident.household if resident else None
    low_income_record = get_latest_low_income_record(resident) if resident else None
    return {
        'resident': resident,
        'household': household,
        'low_income_record': low_income_record,
        'full_name': serialize_scalar(mapped_row.get('full_name')) or (resident.full_name if resident else ''),
        'identity_number': identity_number,
        'job_name': serialize_scalar(mapped_row.get('job_name')),
        'department': serialize_scalar(mapped_row.get('department')),
        'start_date': parse_date_value(mapped_row.get('start_date')),
        'end_date': parse_date_value(mapped_row.get('end_date')),
        'subsidy_amount': parse_decimal_value(mapped_row.get('subsidy_amount')),
        'required_attendance_days': parse_int_value(mapped_row.get('required_attendance_days')),
        'actual_attendance_days': parse_int_value(mapped_row.get('actual_attendance_days')),
        'status': serialize_scalar(mapped_row.get('status')) or PublicJobRecord.STATUS_ON_DUTY,
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_public_job_normalized(normalized, original_row):
    errors = []
    id_valid, id_message = validate_identity_number(normalized['identity_number'])
    if not id_valid:
        errors.append(id_message)
    if not normalized['full_name']:
        errors.append('姓名不能为空')
    if not normalized['job_name']:
        errors.append('岗位名称不能为空')
    if normalized['status'] not in PUBLIC_JOB_STATUS_OPTIONS:
        errors.append('状态必须为在岗、离岗或待岗')
    if serialize_scalar(original_row.get('subsidy_amount')) and normalized['subsidy_amount'] is None:
        errors.append('月补贴标准必须是数字')
    if normalized['end_date'] and normalized['start_date'] and normalized['end_date'] < normalized['start_date']:
        errors.append('合同结束日期不能早于开始日期')
    return errors


def save_public_job_from_normalized(normalized):
    resident = normalized['resident']
    household = normalized['household']
    low_income_record = normalized['low_income_record']
    defaults = {
        'resident': resident,
        'household': household,
        'low_income_record': low_income_record,
        'full_name': normalized['full_name'],
        'gender': resident.gender if resident else '',
        'phone': resident.phone if resident else '',
        'head_name': household.head_name if household else '',
        'village_group': resident.village_group if resident else (household.village_group if household else ''),
        'household_type': resident.household_type if resident else (household.household_type if household else ''),
        'low_income_type': low_income_record.policy_type if low_income_record else '',
        'department': normalized['department'],
        'start_date': normalized['start_date'],
        'end_date': normalized['end_date'],
        'subsidy_amount': normalized['subsidy_amount'],
        'required_attendance_days': normalized['required_attendance_days'],
        'actual_attendance_days': normalized['actual_attendance_days'],
        'status': normalized['status'],
        'notes': normalized['notes'],
    }
    item, created = PublicJobRecord.objects.update_or_create(
        identity_number=normalized['identity_number'],
        job_name=normalized['job_name'],
        start_date=normalized['start_date'],
        defaults=defaults,
    )
    return item, created


def build_public_job_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, PUBLIC_JOB_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_public_job_payload(mapped)
        errors = validate_public_job_normalized(normalized, mapped)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'job_name': normalized['job_name'],
                    'department': normalized['department'],
                    'status': normalized['status'],
                    'subsidy_amount': str(normalized['subsidy_amount']) if normalized['subsidy_amount'] is not None else '',
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = PublicJobImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_public_job_import(batch, mapping):
    build_public_job_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, PUBLIC_JOB_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_public_job_payload(mapped)
        errors = validate_public_job_normalized(normalized, mapped)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_public_job_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = PublicJobImportBatch.STATUS_IMPORTED if not skipped else PublicJobImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


PUBLIC_JOB_EXPORT_COLUMNS = [
    ('姓名', 'full_name'),
    ('身份证号', 'identity_number'),
    ('性别', 'gender'),
    ('村组', 'village_group'),
    ('户主', 'head_name'),
    ('户属性', 'household_type'),
    ('低收入类型', 'low_income_type'),
    ('岗位名称', 'job_name'),
    ('主管部门', 'department'),
    ('合同期', 'contract_period'),
    ('月补贴标准', 'subsidy_amount'),
    ('规定出勤天数', 'required_attendance_days'),
    ('实际出勤天数', 'actual_attendance_days'),
    ('状态', 'status'),
    ('备注', 'notes'),
]


def build_public_job_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '公益性岗位台账'
    data = list_public_job_records({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(PUBLIC_JOB_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(PUBLIC_JOB_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def serialize_care_object(item, index=None):
    resident = item.resident
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'full_name': item.full_name,
        'identity_number': item.identity_number,
        'gender': item.gender,
        'ethnicity': item.ethnicity,
        'age': resident.age if resident else None,
        'phone': item.phone,
        'village_group': item.village_group,
        'address': item.address,
        'care_type': item.care_type,
        'care_level': item.care_level,
        'caregiver_name': item.caregiver_name,
        'caregiver_phone': item.caregiver_phone,
        'notes': item.notes,
    }


def list_care_objects(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = CareObject.objects.select_related('resident', 'household').all()
    if params.get('full_name'):
        queryset = queryset.filter(full_name__icontains=params['full_name'])
    if params.get('identity_number'):
        queryset = queryset.filter(identity_number__icontains=params['identity_number'])
    if params.get('care_type'):
        queryset = queryset.filter(care_type=params['care_type'])
    if params.get('village_group'):
        queryset = queryset.filter(village_group=params['village_group'])
    if params.get('resident_id'):
        queryset = queryset.filter(resident_id=params['resident_id'])

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_care_object(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'filter_options': {
            'care_types': list(
                CareObject.objects.exclude(care_type='').values_list('care_type', flat=True).distinct().order_by('care_type')
            ),
            'village_groups': list(
                CareObject.objects.exclude(village_group='').values_list('village_group', flat=True).distinct().order_by('village_group')
            ),
        },
    }


@transaction.atomic
def create_care_object(data):
    resident_id = data.get('resident_id')
    if not resident_id:
        raise ValueError('必须先选择居民')
    resident = Resident.objects.select_related('household').get(id=resident_id)
    household = resident.household
    item, _ = CareObject.objects.update_or_create(
        resident=resident,
        defaults={
            'household': household,
            'full_name': resident.full_name,
            'identity_number': resident.identity_number,
            'gender': resident.gender,
            'ethnicity': resident.ethnicity,
            'phone': resident.phone,
            'village_group': resident.village_group or (household.village_group if household else ''),
            'address': resident.address or (household.address if household else ''),
            'care_type': serialize_scalar(data.get('care_type')),
            'care_level': serialize_scalar(data.get('care_level')),
            'caregiver_name': serialize_scalar(data.get('caregiver_name')),
            'caregiver_phone': serialize_scalar(data.get('caregiver_phone')),
            'notes': serialize_scalar(data.get('notes')),
        },
    )
    return serialize_care_object(item)


def get_care_object_detail(record_id):
    item = CareObject.objects.select_related('resident', 'household').get(id=record_id)
    return serialize_care_object(item)


@transaction.atomic
def update_care_object(record_id, data):
    item = CareObject.objects.select_related('resident', 'household').get(id=record_id)
    if 'resident_id' in data and data['resident_id'] != item.resident_id:
        resident = Resident.objects.select_related('household').get(id=data['resident_id'])
        household = resident.household
        item.resident = resident
        item.household = household
        item.full_name = resident.full_name
        item.identity_number = resident.identity_number
        item.gender = resident.gender
        item.ethnicity = resident.ethnicity
        item.phone = resident.phone
        item.village_group = resident.village_group or (household.village_group if household else '')
        item.address = resident.address or (household.address if household else '')
    if 'care_type' in data:
        item.care_type = serialize_scalar(data.get('care_type'))
    if 'care_level' in data:
        item.care_level = serialize_scalar(data.get('care_level'))
    if 'caregiver_name' in data:
        item.caregiver_name = serialize_scalar(data.get('caregiver_name'))
    if 'caregiver_phone' in data:
        item.caregiver_phone = serialize_scalar(data.get('caregiver_phone'))
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    item.save()
    return serialize_care_object(item)


@transaction.atomic
def delete_care_object(record_id):
    item = CareObject.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def normalize_care_object_payload(mapped_row):
    identity_number = normalize_identity_number(mapped_row.get('identity_number'))
    resident = Resident.objects.select_related('household').filter(identity_number=identity_number).first() if identity_number else None
    household = resident.household if resident else None
    return {
        'resident': resident,
        'household': household,
        'full_name': serialize_scalar(mapped_row.get('full_name')) or (resident.full_name if resident else ''),
        'identity_number': identity_number,
        'gender': serialize_scalar(mapped_row.get('gender')) or (resident.gender if resident else ''),
        'ethnicity': serialize_scalar(mapped_row.get('ethnicity')) or (resident.ethnicity if resident else ''),
        'phone': serialize_scalar(mapped_row.get('phone')) or (resident.phone if resident else ''),
        'village_group': serialize_scalar(mapped_row.get('village_group')) or (resident.village_group if resident else ''),
        'address': serialize_scalar(mapped_row.get('address')) or (resident.address if resident else ''),
        'care_type': serialize_scalar(mapped_row.get('care_type')),
        'care_level': serialize_scalar(mapped_row.get('care_level')),
        'caregiver_name': serialize_scalar(mapped_row.get('caregiver_name')),
        'caregiver_phone': serialize_scalar(mapped_row.get('caregiver_phone')),
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_care_object_normalized(normalized):
    errors = []
    id_valid, id_message = validate_identity_number(normalized['identity_number'])
    if not id_valid:
        errors.append(id_message)
    if not normalized['full_name']:
        errors.append('姓名不能为空')
    return errors


def save_care_object_from_normalized(normalized):
    item, created = CareObject.objects.update_or_create(
        identity_number=normalized['identity_number'],
        defaults={
            'resident': normalized['resident'],
            'household': normalized['household'],
            'full_name': normalized['full_name'],
            'gender': normalized['gender'],
            'ethnicity': normalized['ethnicity'],
            'phone': normalized['phone'],
            'village_group': normalized['village_group'],
            'address': normalized['address'],
            'care_type': normalized['care_type'],
            'care_level': normalized['care_level'],
            'caregiver_name': normalized['caregiver_name'],
            'caregiver_phone': normalized['caregiver_phone'],
            'notes': normalized['notes'],
        },
    )
    return item, created


def build_care_object_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, CARE_OBJECT_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_care_object_payload(mapped)
        errors = validate_care_object_normalized(normalized)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'care_type': normalized['care_type'],
                    'care_level': normalized['care_level'],
                    'caregiver_name': normalized['caregiver_name'],
                    'caregiver_phone': normalized['caregiver_phone'],
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = CareObjectImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_care_object_import(batch, mapping):
    build_care_object_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, CARE_OBJECT_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_care_object_payload(mapped)
        errors = validate_care_object_normalized(normalized)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_care_object_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'full_name': normalized['full_name'],
                    'identity_number': normalized['identity_number'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = CareObjectImportBatch.STATUS_IMPORTED if not skipped else CareObjectImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


CARE_OBJECT_EXPORT_COLUMNS = [
    ('姓名', 'full_name'),
    ('性别', 'gender'),
    ('民族', 'ethnicity'),
    ('身份证号', 'identity_number'),
    ('年龄', 'age'),
    ('联系电话', 'phone'),
    ('村组', 'village_group'),
    ('家庭地址', 'address'),
    ('关爱类型', 'care_type'),
    ('关爱等级', 'care_level'),
    ('关爱人员', 'caregiver_name'),
    ('联系方式', 'caregiver_phone'),
    ('备注', 'notes'),
]


def build_care_object_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '关爱对象台账'
    data = list_care_objects({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(CARE_OBJECT_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(CARE_OBJECT_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


MEDIATION_DISPUTE_TYPE_OPTIONS = [
    '婚姻家庭纠纷',
    '邻里纠纷',
    '土地承包纠纷',
    '债务纠纷',
    '劳动争议',
]
MEDIATION_STATUS_OPTIONS = {choice[0] for choice in MediationRecord.STATUS_CHOICES}
MEDIATION_EXPORT_COLUMNS = [
    ('档案编号', 'archive_number'),
    ('纠纷类型', 'dispute_type'),
    ('申请人', 'applicant_names'),
    ('被申请人', 'respondent_names'),
    ('状态', 'status'),
    ('申请日期', 'application_date'),
    ('发生日期', 'occurrence_date'),
    ('发生地点', 'occurrence_location'),
    ('纠纷概述', 'dispute_summary'),
    ('申请事项', 'application_request_text'),
    ('创建时间', 'created_at'),
]


def format_display_date(value):
    if not value:
        return ''
    if hasattr(value, 'strftime'):
        return value.strftime('%Y-%m-%d')
    return str(value)


def generate_next_mediation_archive_number():
    max_index = 0
    for archive_number in MediationRecord.objects.exclude(archive_number='').values_list('archive_number', flat=True):
        match = re.search(r'(\d+)$', archive_number)
        if match:
            max_index = max(max_index, int(match.group(1)))
    return f'TJ-{max_index + 1:03d}'


def normalize_mediation_party(party):
    if not isinstance(party, dict):
        party = {}
    identity_number = normalize_identity_number(party.get('identity_number'))
    age = None
    age_text = serialize_scalar(party.get('age'))
    if age_text:
        try:
            age = int(float(age_text))
        except (TypeError, ValueError):
            age = None
    if age is None and identity_number and len(identity_number) >= 14:
        try:
            birth_date = datetime.strptime(identity_number[6:14], '%Y%m%d').date()
            age = calculate_age(birth_date)
        except ValueError:
            age = None
    return {
        'name': serialize_scalar(party.get('name')),
        'gender': serialize_scalar(party.get('gender')),
        'ethnicity': serialize_scalar(party.get('ethnicity')),
        'age': age,
        'identity_number': identity_number,
        'phone': serialize_scalar(party.get('phone')),
        'occupation': serialize_scalar(party.get('occupation')),
        'address': serialize_scalar(party.get('address')),
    }


def normalize_mediation_party_list(parties):
    normalized_parties = []
    for party in parties or []:
        normalized = normalize_mediation_party(party)
        if any(value not in ('', None) for value in normalized.values()):
            normalized_parties.append(normalized)
    return normalized_parties


def validate_mediation_party_list(parties, label):
    errors = []
    if not parties:
        errors.append(f'请至少填写一位{label}')
        return errors
    for index, party in enumerate(parties, start=1):
        if not party['name']:
            errors.append(f'{label}{index}姓名不能为空')
        if party['identity_number']:
            id_valid, id_message = validate_identity_number(party['identity_number'])
            if not id_valid:
                errors.append(f'{label}{index}{id_message}')
    return errors


def normalize_application_requests(value):
    if isinstance(value, list):
        lines = [serialize_scalar(item) for item in value]
    else:
        lines = str(value or '').splitlines()
    return [line.strip() for line in lines if str(line).strip()]


def join_party_names(parties):
    names = [party.get('name', '') for party in parties if party.get('name')]
    return '、'.join(names)


def serialize_mediation_record(item, index=None):
    applicants = normalize_mediation_party_list(item.applicants)
    respondents = normalize_mediation_party_list(item.respondents)
    application_requests = normalize_application_requests(item.application_requests)
    return {
        'id': item.id,
        'seq': index,
        'archive_number': item.archive_number,
        'dispute_type': item.dispute_type,
        'status': item.status,
        'application_date': format_display_date(item.application_date),
        'occurrence_date': format_display_date(item.occurrence_date),
        'occurrence_location': item.occurrence_location,
        'applicants': applicants,
        'respondents': respondents,
        'applicant_names': join_party_names(applicants),
        'respondent_names': join_party_names(respondents),
        'dispute_summary': item.dispute_summary,
        'application_requests': application_requests,
        'application_request_text': '\n'.join(application_requests),
        'created_at': format_display_date(timezone.localtime(item.created_at)),
        'updated_at': format_display_date(timezone.localtime(item.updated_at)),
    }


def list_mediation_records(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = MediationRecord.objects.all()
    if params.get('archive_number'):
        queryset = queryset.filter(archive_number__icontains=params['archive_number'])
    if params.get('dispute_type'):
        queryset = queryset.filter(dispute_type=params['dispute_type'])
    if params.get('status'):
        queryset = queryset.filter(status=params['status'])
    created_from = parse_date_value(params.get('created_from'))
    created_to = parse_date_value(params.get('created_to'))
    if created_from:
        queryset = queryset.filter(created_at__date__gte=created_from)
    if created_to:
        queryset = queryset.filter(created_at__date__lte=created_to)

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_mediation_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'filter_options': {
            'dispute_types': MEDIATION_DISPUTE_TYPE_OPTIONS,
            'statuses': [choice[0] for choice in MediationRecord.STATUS_CHOICES],
        },
    }


def validate_mediation_payload(data):
    applicants = normalize_mediation_party_list(data.get('applicants') or [])
    respondents = normalize_mediation_party_list(data.get('respondents') or [])
    dispute_type = serialize_scalar(data.get('dispute_type'))
    application_requests = normalize_application_requests(data.get('application_requests'))
    status = serialize_scalar(data.get('status')) or MediationRecord.STATUS_IN_PROGRESS
    errors = []
    if not dispute_type:
        errors.append('请选择纠纷类型')
    errors.extend(validate_mediation_party_list(applicants, '申请人'))
    errors.extend(validate_mediation_party_list(respondents, '被申请人'))
    if not application_requests:
        errors.append('请至少填写一条申请事项')
    if status not in MEDIATION_STATUS_OPTIONS:
        errors.append('状态不合法')
    if errors:
        raise ValueError('；'.join(errors))
    return {
        'archive_number': serialize_scalar(data.get('archive_number')),
        'dispute_type': dispute_type,
        'status': status,
        'application_date': parse_date_value(data.get('application_date')),
        'occurrence_date': parse_date_value(data.get('occurrence_date')),
        'occurrence_location': serialize_scalar(data.get('occurrence_location')),
        'applicants': applicants,
        'respondents': respondents,
        'dispute_summary': serialize_scalar(data.get('dispute_summary')),
        'application_requests': application_requests,
    }


@transaction.atomic
def create_mediation_record(data):
    normalized = validate_mediation_payload(data)
    archive_number = normalized['archive_number'] or generate_next_mediation_archive_number()
    if MediationRecord.objects.filter(archive_number=archive_number).exists():
        archive_number = generate_next_mediation_archive_number()
    item = MediationRecord.objects.create(
        archive_number=archive_number,
        dispute_type=normalized['dispute_type'],
        status=normalized['status'],
        application_date=normalized['application_date'],
        occurrence_date=normalized['occurrence_date'],
        occurrence_location=normalized['occurrence_location'],
        applicants=normalized['applicants'],
        respondents=normalized['respondents'],
        dispute_summary=normalized['dispute_summary'],
        application_requests=normalized['application_requests'],
    )
    return serialize_mediation_record(item)


def get_mediation_record_detail(record_id):
    item = MediationRecord.objects.get(id=record_id)
    return serialize_mediation_record(item)


@transaction.atomic
def update_mediation_record(record_id, data):
    item = MediationRecord.objects.get(id=record_id)
    normalized = validate_mediation_payload(data)
    item.dispute_type = normalized['dispute_type']
    item.status = normalized['status']
    item.application_date = normalized['application_date']
    item.occurrence_date = normalized['occurrence_date']
    item.occurrence_location = normalized['occurrence_location']
    item.applicants = normalized['applicants']
    item.respondents = normalized['respondents']
    item.dispute_summary = normalized['dispute_summary']
    item.application_requests = normalized['application_requests']
    item.save()
    return serialize_mediation_record(item)


@transaction.atomic
def delete_mediation_record(record_id):
    item = MediationRecord.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def build_mediation_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '人民调解台账'
    data = list_mediation_records({**params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(MEDIATION_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(MEDIATION_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def set_document_default_style(document):
    style = document.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)


def add_document_line(document, text='', bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=12):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    return paragraph


def fill_document_text(value, placeholder=''):
    return serialize_scalar(value) or placeholder


def build_mediation_application_document(record):
    payload = serialize_mediation_record(record)
    applicants = payload['applicants'] or [{}]
    respondents = payload['respondents'] or [{}]
    applicant = applicants[0]
    respondent = respondents[0]

    document = Document()
    set_document_default_style(document)
    add_document_line(document, '人民调解申请书', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=18)
    add_document_line(document, '')
    add_document_line(
        document,
        f"申请人姓名 {fill_document_text(applicant.get('name'), '        ')}    性别 {fill_document_text(applicant.get('gender'), '    ')}    民族 {fill_document_text(applicant.get('ethnicity'), '    ')}    年龄 {fill_document_text(applicant.get('age'), '    ')}",
    )
    add_document_line(
        document,
        f"职业或职务 {fill_document_text(applicant.get('occupation'), '                ')}    联系方式 {fill_document_text(applicant.get('phone'), '                ')}",
    )
    add_document_line(document, f"单位或住址 {fill_document_text(applicant.get('address'), '                                        ')}")
    add_document_line(document, '')
    add_document_line(
        document,
        f"被申请人姓名 {fill_document_text(respondent.get('name'), '        ')}    性别 {fill_document_text(respondent.get('gender'), '    ')}    民族 {fill_document_text(respondent.get('ethnicity'), '    ')}    年龄 {fill_document_text(respondent.get('age'), '    ')}",
    )
    add_document_line(
        document,
        f"职业或职务 {fill_document_text(respondent.get('occupation'), '                ')}    联系方式 {fill_document_text(respondent.get('phone'), '                ')}",
    )
    add_document_line(document, f"单位或住址 {fill_document_text(respondent.get('address'), '                                        ')}")
    add_document_line(document, '')
    add_document_line(document, '纠纷简要情况：')
    summary_lines = payload['dispute_summary'].splitlines() if payload['dispute_summary'] else []
    for line in (summary_lines[:6] or ['']):
        add_document_line(document, fill_document_text(line, ''))
    add_document_line(document, '')
    add_document_line(document, '当事人申请事项：')
    request_lines = payload['application_requests'][:3]
    while len(request_lines) < 3:
        request_lines.append('')
    for index, line in enumerate(request_lines, start=1):
        add_document_line(document, f'{index}. {fill_document_text(line, "")}')
    add_document_line(document, '')
    add_document_line(document, '人民调解委员会已将申请人民调解的相关规定告知我，')
    add_document_line(document, '现自愿申请人民调解委员会进行调解。')
    add_document_line(document, '')
    add_document_line(document, '申请人（签名盖章或按指印）：')
    sign_date = record.application_date or timezone.localdate()
    add_document_line(document, f'{sign_date.year} 年 {sign_date.month} 月 {sign_date.day} 日', align=WD_ALIGN_PARAGRAPH.RIGHT)
    return document


PROJECT_SOURCE_OPTIONS = {choice[0] for choice in ProjectRecord.SOURCE_CHOICES}
PROJECT_TYPE_OPTIONS = {choice[0] for choice in ProjectRecord.TYPE_CHOICES}
PROJECT_STATUS_OPTIONS = {choice[0] for choice in ProjectRecord.STATUS_CHOICES}


def serialize_project_record(item, index=None):
    return {
        'id': item.id,
        'seq': index,
        'project_name': item.project_name,
        'project_source': item.project_source,
        'project_type': item.project_type,
        'secondary_type': item.secondary_type,
        'project_status': item.project_status,
        'planning_year': item.planning_year,
        'implementation_year': item.implementation_year,
        'included_in_plan': item.included_in_plan,
        'included_in_plan_label': '是' if item.included_in_plan else '否',
        'total_investment': str(item.total_investment) if item.total_investment is not None else '',
        'settled_amount': str(item.settled_amount) if item.settled_amount is not None else '',
        'audited_amount': str(item.audited_amount) if item.audited_amount is not None else '',
        'responsible_person': item.responsible_person,
        'project_location': item.project_location,
        'project_description': item.project_description,
        'notes': item.notes,
    }


def list_projects(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = ProjectRecord.objects.all()

    keyword = serialize_scalar(params.get('keyword'))
    if keyword:
        queryset = queryset.filter(
            Q(project_name__icontains=keyword)
            | Q(project_location__icontains=keyword)
            | Q(project_description__icontains=keyword)
            | Q(responsible_person__icontains=keyword)
        )
    if params.get('project_source'):
        queryset = queryset.filter(project_source=params['project_source'])
    if params.get('project_type'):
        queryset = queryset.filter(project_type=params['project_type'])
    if params.get('project_status'):
        queryset = queryset.filter(project_status=params['project_status'])
    if params.get('included_in_plan') not in (None, '', '全部'):
        queryset = queryset.filter(included_in_plan=parse_bool_value(params.get('included_in_plan')))

    planning_year_start = parse_int_value(params.get('planning_year_start'), None)
    planning_year_end = parse_int_value(params.get('planning_year_end'), None)
    implementation_year_start = parse_int_value(params.get('implementation_year_start'), None)
    implementation_year_end = parse_int_value(params.get('implementation_year_end'), None)
    if planning_year_start is not None:
        queryset = queryset.filter(planning_year__gte=planning_year_start)
    if planning_year_end is not None:
        queryset = queryset.filter(planning_year__lte=planning_year_end)
    if implementation_year_start is not None:
        queryset = queryset.filter(implementation_year__gte=implementation_year_start)
    if implementation_year_end is not None:
        queryset = queryset.filter(implementation_year__lte=implementation_year_end)

    stats_queryset = list(queryset.values_list('total_investment', 'settled_amount', 'audited_amount'))
    total_investment = sum((amount or Decimal('0')) for amount, _, _ in stats_queryset)
    settled_amount = sum((amount or Decimal('0')) for _, amount, _ in stats_queryset)
    audited_amount = sum((amount or Decimal('0')) for _, _, amount in stats_queryset)

    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_project_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'stats': {
            'project_count': paginator.count,
            'total_investment': str(total_investment),
            'settled_amount': str(settled_amount),
            'audited_amount': str(audited_amount),
        },
    }


@transaction.atomic
def create_project_record(data):
    project_name = serialize_scalar(data.get('project_name'))
    if not project_name:
        raise ValueError('项目名称不能为空')
    item = ProjectRecord.objects.create(
        project_name=project_name,
        project_source=serialize_scalar(data.get('project_source')),
        project_type=serialize_scalar(data.get('project_type')),
        secondary_type=serialize_scalar(data.get('secondary_type')),
        project_status=serialize_scalar(data.get('project_status')) or '规划中',
        planning_year=parse_int_value(data.get('planning_year'), None),
        implementation_year=parse_int_value(data.get('implementation_year'), None),
        included_in_plan=parse_bool_value(data.get('included_in_plan')),
        total_investment=parse_decimal_value(data.get('total_investment')),
        settled_amount=parse_decimal_value(data.get('settled_amount')),
        audited_amount=parse_decimal_value(data.get('audited_amount')),
        responsible_person=serialize_scalar(data.get('responsible_person')),
        project_location=serialize_scalar(data.get('project_location')),
        project_description=serialize_scalar(data.get('project_description')),
        notes=serialize_scalar(data.get('notes')),
    )
    return serialize_project_record(item)


def get_project_record_detail(record_id):
    item = ProjectRecord.objects.get(id=record_id)
    return serialize_project_record(item)


@transaction.atomic
def update_project_record(record_id, data):
    item = ProjectRecord.objects.get(id=record_id)
    if 'project_name' in data:
        project_name = serialize_scalar(data.get('project_name'))
        if not project_name:
            raise ValueError('项目名称不能为空')
        item.project_name = project_name
    if 'project_source' in data:
        item.project_source = serialize_scalar(data.get('project_source'))
    if 'project_type' in data:
        item.project_type = serialize_scalar(data.get('project_type'))
    if 'secondary_type' in data:
        item.secondary_type = serialize_scalar(data.get('secondary_type'))
    if 'project_status' in data:
        item.project_status = serialize_scalar(data.get('project_status')) or '规划中'
    if 'planning_year' in data:
        item.planning_year = parse_int_value(data.get('planning_year'), None)
    if 'implementation_year' in data:
        item.implementation_year = parse_int_value(data.get('implementation_year'), None)
    if 'included_in_plan' in data:
        item.included_in_plan = parse_bool_value(data.get('included_in_plan'))
    if 'total_investment' in data:
        item.total_investment = parse_decimal_value(data.get('total_investment'))
    if 'settled_amount' in data:
        item.settled_amount = parse_decimal_value(data.get('settled_amount'))
    if 'audited_amount' in data:
        item.audited_amount = parse_decimal_value(data.get('audited_amount'))
    if 'responsible_person' in data:
        item.responsible_person = serialize_scalar(data.get('responsible_person'))
    if 'project_location' in data:
        item.project_location = serialize_scalar(data.get('project_location'))
    if 'project_description' in data:
        item.project_description = serialize_scalar(data.get('project_description'))
    if 'notes' in data:
        item.notes = serialize_scalar(data.get('notes'))
    item.save()
    return serialize_project_record(item)


@transaction.atomic
def delete_project_record(record_id):
    item = ProjectRecord.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


def normalize_project_payload(mapped_row):
    return {
        'project_name': serialize_scalar(mapped_row.get('project_name')),
        'project_source': serialize_scalar(mapped_row.get('project_source')),
        'project_type': serialize_scalar(mapped_row.get('project_type')),
        'secondary_type': serialize_scalar(mapped_row.get('secondary_type')),
        'project_status': serialize_scalar(mapped_row.get('project_status')) or '规划中',
        'planning_year': parse_int_value(mapped_row.get('planning_year'), None),
        'implementation_year': parse_int_value(mapped_row.get('implementation_year'), None),
        'included_in_plan': parse_bool_value(mapped_row.get('included_in_plan')),
        'total_investment': parse_decimal_value(mapped_row.get('total_investment')),
        'settled_amount': parse_decimal_value(mapped_row.get('settled_amount')),
        'audited_amount': parse_decimal_value(mapped_row.get('audited_amount')),
        'responsible_person': serialize_scalar(mapped_row.get('responsible_person')),
        'project_location': serialize_scalar(mapped_row.get('project_location')),
        'project_description': serialize_scalar(mapped_row.get('project_description')),
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_project_normalized(normalized, original_row):
    errors = []
    if not normalized['project_name']:
        errors.append('项目名称不能为空')
    if normalized['project_source'] and normalized['project_source'] not in PROJECT_SOURCE_OPTIONS:
        errors.append('项目库来源不在允许范围内')
    if normalized['project_type'] and normalized['project_type'] not in PROJECT_TYPE_OPTIONS:
        errors.append('项目类型必须为基础设施、产业发展、社会事业或生态环境')
    if normalized['project_status'] and normalized['project_status'] not in PROJECT_STATUS_OPTIONS:
        errors.append('项目状态必须为规划中、实施中、已完成或已终止')
    if normalized['planning_year'] is not None and normalized['planning_year'] <= 0:
        errors.append('规划年度必须为正整数')
    if normalized['implementation_year'] is not None and normalized['implementation_year'] <= 0:
        errors.append('实施年度必须为正整数')
    if serialize_scalar(original_row.get('total_investment')) and normalized['total_investment'] is None:
        errors.append('项目预算总投资必须是数字')
    if serialize_scalar(original_row.get('settled_amount')) and normalized['settled_amount'] is None:
        errors.append('结算金额必须是数字')
    if serialize_scalar(original_row.get('audited_amount')) and normalized['audited_amount'] is None:
        errors.append('决算审计金额必须是数字')
    return errors


def save_project_from_normalized(normalized):
    item, created = ProjectRecord.objects.update_or_create(
        project_name=normalized['project_name'],
        project_source=normalized['project_source'],
        planning_year=normalized['planning_year'],
        implementation_year=normalized['implementation_year'],
        defaults={
            'project_type': normalized['project_type'],
            'secondary_type': normalized['secondary_type'],
            'project_status': normalized['project_status'],
            'included_in_plan': normalized['included_in_plan'],
            'total_investment': normalized['total_investment'],
            'settled_amount': normalized['settled_amount'],
            'audited_amount': normalized['audited_amount'],
            'responsible_person': normalized['responsible_person'],
            'project_location': normalized['project_location'],
            'project_description': normalized['project_description'],
            'notes': normalized['notes'],
        },
    )
    return item, created


def build_project_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, PROJECT_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_project_payload(mapped)
        errors = validate_project_normalized(normalized, mapped)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'project_name': normalized['project_name'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'project_name': normalized['project_name'],
                    'project_type': normalized['project_type'],
                    'project_status': normalized['project_status'],
                    'implementation_year': normalized['implementation_year'] or '',
                    'total_investment': str(normalized['total_investment']) if normalized['total_investment'] is not None else '',
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = ProjectImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_project_import(batch, mapping):
    build_project_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, PROJECT_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_project_payload(mapped)
        errors = validate_project_normalized(normalized, mapped)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'project_name': normalized['project_name'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_project_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'project_name': normalized['project_name'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = ProjectImportBatch.STATUS_IMPORTED if not skipped else ProjectImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


PROJECT_EXPORT_COLUMNS = [
    ('项目名称', 'project_name'),
    ('项目库来源', 'project_source'),
    ('项目类型', 'project_type'),
    ('二级类型', 'secondary_type'),
    ('项目状态', 'project_status'),
    ('规划年度', 'planning_year'),
    ('实施年度', 'implementation_year'),
    ('纳入计划', 'included_in_plan_label'),
    ('项目预算总投资(万元)', 'total_investment'),
    ('结算金额(万元)', 'settled_amount'),
    ('决算审计金额(万元)', 'audited_amount'),
    ('督护人/责任人', 'responsible_person'),
    ('项目地点', 'project_location'),
    ('项目描述', 'project_description'),
]


def build_project_export_workbook(params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = '项目台账'
    normalized_params = {key: params.get(key) for key in params} if hasattr(params, 'get') else dict(params)
    data = list_projects({**normalized_params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(PROJECT_EXPORT_COLUMNS, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(PROJECT_EXPORT_COLUMNS, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


def is_farmland_transfered(value):
    text = serialize_scalar(value)
    if not text:
        return False
    return text not in {'未流转', '否', '无', '未', '正常'}


def is_farmland_confirmed(value):
    text = serialize_scalar(value)
    if not text:
        return False
    return text in {'已确权', '是', '已完成', '已办理', '完成'} or '已确权' in text


def is_farmland_reclaimed(item):
    plot_status = serialize_scalar(item.plot_status)
    latest_change = serialize_scalar(item.latest_change)
    return '复垦' in plot_status or '复垦' in latest_change


def serialize_farmland_record(item, index=None):
    resident = item.resident
    return {
        'id': item.id,
        'seq': index,
        'resident_id': item.resident_id,
        'household_id': item.household_id,
        'linked_resident_id': item.linked_resident_id,
        'resident_name': resident.full_name if resident else '',
        'plot_code': item.plot_code,
        'village_group': item.village_group,
        'contractor_name': item.contractor_name,
        'contractor_identity_number': item.contractor_identity_number,
        'plot_location': item.plot_location,
        'area_mu': str(item.area_mu) if item.area_mu is not None else '',
        'east_boundary': item.east_boundary,
        'south_boundary': item.south_boundary,
        'west_boundary': item.west_boundary,
        'north_boundary': item.north_boundary,
        'boundaries': '；'.join(
            filter(
                None,
                [
                    f'东至{item.east_boundary}' if item.east_boundary else '',
                    f'南至{item.south_boundary}' if item.south_boundary else '',
                    f'西至{item.west_boundary}' if item.west_boundary else '',
                    f'北至{item.north_boundary}' if item.north_boundary else '',
                ],
            )
        ),
        'plot_status': item.plot_status,
        'transfer_status': item.transfer_status,
        'confirmation_status': item.confirmation_status,
        'current_planting': item.current_planting,
        'latest_change': item.latest_change,
        'change_date': item.change_date.isoformat() if item.change_date else '',
        'notes': item.notes,
    }


def apply_farmland_filters(queryset, params):
    contractor_name = serialize_scalar(params.get('contractor_name'))
    if contractor_name:
        queryset = queryset.filter(contractor_name__icontains=contractor_name)
    plot_code = serialize_scalar(params.get('plot_code'))
    if plot_code:
        queryset = queryset.filter(plot_code__icontains=plot_code)
    village_group = serialize_scalar(params.get('village_group'))
    if village_group and village_group != '全部村组':
        queryset = queryset.filter(village_group=village_group)
    plot_status = serialize_scalar(params.get('plot_status'))
    if plot_status and plot_status != '全部状态':
        queryset = queryset.filter(plot_status=plot_status)
    return queryset


def build_farmland_stats(queryset):
    rows = list(queryset.values_list('area_mu', 'plot_status', 'transfer_status', 'latest_change'))
    total_area = sum((area or Decimal('0')) for area, _, _, _ in rows)
    return {
        'plot_count': len(rows),
        'total_area_mu': str(total_area),
        'transferred_count': sum(1 for _, _, transfer_status, _ in rows if is_farmland_transfered(transfer_status)),
        'abandoned_count': sum(1 for _, plot_status, _, _ in rows if '撂荒' in serialize_scalar(plot_status)),
        'reclaimed_count': sum(
            1
            for _, plot_status, _, latest_change in rows
            if '复垦' in serialize_scalar(plot_status) or '复垦' in serialize_scalar(latest_change)
        ),
    }


def get_farmland_filter_options():
    return {
        'village_groups': list(
            FarmlandRecord.objects.exclude(village_group='').order_by('village_group').values_list('village_group', flat=True).distinct()
        ),
        'plot_statuses': list(
            FarmlandRecord.objects.exclude(plot_status='').order_by('plot_status').values_list('plot_status', flat=True).distinct()
        ),
    }


def list_farmland_records(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = apply_farmland_filters(FarmlandRecord.objects.select_related('resident', 'household').all(), params)
    stats = build_farmland_stats(queryset)
    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_farmland_record(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'stats': stats,
        'filter_options': get_farmland_filter_options(),
    }


def list_farmland_households(params):
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = apply_farmland_filters(FarmlandRecord.objects.select_related('resident', 'household').all(), params)
    grouped = {}
    for item in queryset:
        group_key = item.contractor_identity_number or str(item.linked_resident_id or '') or item.contractor_name or item.plot_code
        payload = grouped.setdefault(
            group_key,
            {
                'id': item.id,
                'contractor_name': item.contractor_name,
                'contractor_identity_number': item.contractor_identity_number,
                'linked_resident_id': item.linked_resident_id,
                'village_groups': set(),
                'plot_codes': [],
                'plot_count': 0,
                'total_area_mu': Decimal('0'),
                'transferred_count': 0,
                'confirmed_count': 0,
                'latest_change_date': None,
            },
        )
        if item.village_group:
            payload['village_groups'].add(item.village_group)
        payload['plot_codes'].append(item.plot_code)
        payload['plot_count'] += 1
        payload['total_area_mu'] += item.area_mu or Decimal('0')
        if is_farmland_transfered(item.transfer_status):
            payload['transferred_count'] += 1
        if is_farmland_confirmed(item.confirmation_status):
            payload['confirmed_count'] += 1
        if item.change_date and (payload['latest_change_date'] is None or item.change_date > payload['latest_change_date']):
            payload['latest_change_date'] = item.change_date

    ordered = sorted(
        grouped.values(),
        key=lambda item: (
            '、'.join(sorted(item['village_groups'])),
            item['contractor_name'],
            item['contractor_identity_number'],
        ),
    )
    total = len(ordered)
    start = (page_number - 1) * page_size
    end = start + page_size
    items = []
    for index, item in enumerate(ordered[start:end], start=start + 1):
        items.append(
            {
                'id': item['id'],
                'seq': index,
                'contractor_name': item['contractor_name'],
                'contractor_identity_number': item['contractor_identity_number'],
                'linked_resident_id': item['linked_resident_id'],
                'village_group': '、'.join(sorted(item['village_groups'])),
                'plot_count': item['plot_count'],
                'total_area_mu': str(item['total_area_mu']),
                'transferred_count': item['transferred_count'],
                'confirmed_count': item['confirmed_count'],
                'plot_codes': '、'.join(item['plot_codes']),
                'latest_change_date': item['latest_change_date'].isoformat() if item['latest_change_date'] else '',
            }
        )
    return {
        'items': items,
        'pagination': {
            'page': page_number,
            'page_size': page_size,
            'total': total,
            'total_pages': max(1, (total + page_size - 1) // page_size),
        },
        'stats': build_farmland_stats(queryset),
        'filter_options': get_farmland_filter_options(),
    }


def normalize_farmland_payload(mapped_row):
    linked_resident_id = parse_int_value(mapped_row.get('linked_resident_id'), None)
    resident = None
    if linked_resident_id is not None:
        resident = Resident.objects.select_related('household').filter(id=linked_resident_id).first()
    contractor_identity_number = normalize_identity_number(mapped_row.get('contractor_identity_number'))
    if resident is None and contractor_identity_number:
        resident = Resident.objects.select_related('household').filter(identity_number=contractor_identity_number).first()
    household = resident.household if resident else None
    resident_reference_id = linked_resident_id if linked_resident_id is not None else (resident.id if resident else None)
    return {
        'resident': resident,
        'household': household,
        'plot_code': serialize_scalar(mapped_row.get('plot_code')),
        'village_group': serialize_scalar(mapped_row.get('village_group')) or (resident.village_group if resident else ''),
        'contractor_name': serialize_scalar(mapped_row.get('contractor_name'))
        or (resident.full_name if resident else (household.head_name if household else '')),
        'contractor_identity_number': contractor_identity_number or (resident.identity_number if resident else ''),
        'linked_resident_id': resident_reference_id,
        'plot_location': serialize_scalar(mapped_row.get('plot_location')),
        'area_mu': parse_decimal_value(mapped_row.get('area_mu')),
        'east_boundary': serialize_scalar(mapped_row.get('east_boundary')),
        'south_boundary': serialize_scalar(mapped_row.get('south_boundary')),
        'west_boundary': serialize_scalar(mapped_row.get('west_boundary')),
        'north_boundary': serialize_scalar(mapped_row.get('north_boundary')),
        'plot_status': serialize_scalar(mapped_row.get('plot_status')),
        'transfer_status': serialize_scalar(mapped_row.get('transfer_status')),
        'confirmation_status': serialize_scalar(mapped_row.get('confirmation_status')),
        'current_planting': serialize_scalar(mapped_row.get('current_planting')),
        'latest_change': serialize_scalar(mapped_row.get('latest_change')),
        'change_date': parse_date_value(mapped_row.get('change_date')),
        'notes': serialize_scalar(mapped_row.get('notes')),
    }


def validate_farmland_normalized(normalized, original_row):
    errors = []
    if not normalized['plot_code']:
        errors.append('地块编号不能为空')
    if not normalized['contractor_name']:
        errors.append('承包户不能为空')
    if normalized['contractor_identity_number']:
        id_valid, id_message = validate_identity_number(normalized['contractor_identity_number'])
        if not id_valid:
            errors.append(f'承包户身份证号{ id_message }')
    if serialize_scalar(original_row.get('linked_resident_id')) and normalized['resident'] is None:
        errors.append('关联居民ID不存在')
    if serialize_scalar(original_row.get('area_mu')) and normalized['area_mu'] is None:
        errors.append('面积（亩）必须是数字')
    if serialize_scalar(original_row.get('change_date')) and normalized['change_date'] is None:
        errors.append('变更日期格式不正确')
    return errors


def save_farmland_from_normalized(normalized):
    defaults = {
        'resident': normalized['resident'],
        'household': normalized['household'],
        'village_group': normalized['village_group'],
        'contractor_name': normalized['contractor_name'],
        'contractor_identity_number': normalized['contractor_identity_number'],
        'linked_resident_id': normalized['linked_resident_id'],
        'plot_location': normalized['plot_location'],
        'area_mu': normalized['area_mu'],
        'east_boundary': normalized['east_boundary'],
        'south_boundary': normalized['south_boundary'],
        'west_boundary': normalized['west_boundary'],
        'north_boundary': normalized['north_boundary'],
        'plot_status': normalized['plot_status'],
        'transfer_status': normalized['transfer_status'],
        'confirmation_status': normalized['confirmation_status'],
        'current_planting': normalized['current_planting'],
        'latest_change': normalized['latest_change'],
        'change_date': normalized['change_date'],
        'notes': normalized['notes'],
    }
    item, created = FarmlandRecord.objects.update_or_create(
        plot_code=normalized['plot_code'],
        defaults=defaults,
    )
    return item, created


def build_farmland_preview(batch, mapping, preview_limit=8):
    mapping = normalize_module_mapping(mapping, FARMLAND_SYSTEM_FIELDS)
    if not mapping:
        raise ValueError('请先完成 Excel 字段映射。')
    _, rows = load_excel_rows(batch.file_path)
    preview_rows = []
    valid_count = 0
    invalid_count = 0
    all_errors = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_farmland_payload(mapped)
        errors = validate_farmland_normalized(normalized, mapped)
        if errors:
            invalid_count += 1
            all_errors.append(
                {
                    'row_number': index + 1,
                    'plot_code': normalized['plot_code'],
                    'contractor_name': normalized['contractor_name'],
                    'messages': errors,
                    'original_data': row,
                }
            )
        else:
            valid_count += 1
        if len(preview_rows) < preview_limit:
            preview_rows.append(
                {
                    'row_number': index + 1,
                    'plot_code': normalized['plot_code'],
                    'contractor_name': normalized['contractor_name'],
                    'village_group': normalized['village_group'],
                    'area_mu': str(normalized['area_mu']) if normalized['area_mu'] is not None else '',
                    'plot_status': normalized['plot_status'],
                    'transfer_status': normalized['transfer_status'],
                    'errors': errors,
                }
            )
    batch.field_mapping = mapping
    batch.valid_rows = valid_count
    batch.invalid_rows = invalid_count
    batch.error_details = all_errors
    batch.status = FarmlandImportBatch.STATUS_PREVIEWED
    batch.save(update_fields=['field_mapping', 'valid_rows', 'invalid_rows', 'error_details', 'status', 'updated_at'])
    return {
        'preview_rows': preview_rows,
        'total_rows': len(rows),
        'valid_rows': valid_count,
        'invalid_rows': invalid_count,
        'errors': all_errors[:20],
    }


@transaction.atomic
def commit_farmland_import(batch, mapping):
    build_farmland_preview(batch, mapping, preview_limit=20)
    _, rows = load_excel_rows(batch.file_path)
    mapping = normalize_module_mapping(mapping, FARMLAND_SYSTEM_FIELDS)
    created_count = 0
    updated_count = 0
    skipped = []
    for index, row in enumerate(rows, start=1):
        mapped = map_module_row(row, mapping)
        normalized = normalize_farmland_payload(mapped)
        errors = validate_farmland_normalized(normalized, mapped)
        if errors:
            skipped.append(
                {
                    'row_number': index + 1,
                    'plot_code': normalized['plot_code'],
                    'contractor_name': normalized['contractor_name'],
                    'messages': errors,
                    'original_data': row,
                }
            )
            continue
        try:
            _, created = save_farmland_from_normalized(normalized)
            if created:
                created_count += 1
            else:
                updated_count += 1
        except Exception as exc:
            skipped.append(
                {
                    'row_number': index + 1,
                    'plot_code': normalized['plot_code'],
                    'contractor_name': normalized['contractor_name'],
                    'messages': [f'数据库错误: {exc}'],
                    'original_data': row,
                }
            )
    batch.status = FarmlandImportBatch.STATUS_IMPORTED if not skipped else FarmlandImportBatch.STATUS_FAILED
    batch.imported_rows = created_count + updated_count
    batch.created_rows = created_count
    batch.updated_rows = updated_count
    batch.error_details = skipped
    batch.save(update_fields=['status', 'imported_rows', 'created_rows', 'updated_rows', 'error_details', 'updated_at'])
    return {
        'total_rows': batch.total_rows,
        'valid_rows': batch.valid_rows,
        'invalid_rows': batch.invalid_rows,
        'created_rows': created_count,
        'updated_rows': updated_count,
        'skipped_rows': len(skipped),
        'errors': skipped[:20],
    }


FARMLAND_DETAIL_EXPORT_COLUMNS = [
    ('地块编号', 'plot_code'),
    ('村组', 'village_group'),
    ('承包户', 'contractor_name'),
    ('承包户身份证号', 'contractor_identity_number'),
    ('关联居民ID', 'linked_resident_id'),
    ('地块位置', 'plot_location'),
    ('面积（亩）', 'area_mu'),
    ('东至', 'east_boundary'),
    ('南至', 'south_boundary'),
    ('西至', 'west_boundary'),
    ('北至', 'north_boundary'),
    ('地块状态', 'plot_status'),
    ('流转情况', 'transfer_status'),
    ('确权情况', 'confirmation_status'),
    ('当前种植', 'current_planting'),
    ('最新变更', 'latest_change'),
    ('变更日期', 'change_date'),
    ('备注', 'notes'),
]

FARMLAND_HOUSEHOLD_EXPORT_COLUMNS = [
    ('承包户', 'contractor_name'),
    ('承包户身份证号', 'contractor_identity_number'),
    ('关联居民ID', 'linked_resident_id'),
    ('村组', 'village_group'),
    ('地块数量', 'plot_count'),
    ('总面积（亩）', 'total_area_mu'),
    ('流转地块', 'transferred_count'),
    ('已确权地块', 'confirmed_count'),
    ('地块编号', 'plot_codes'),
    ('最近变更日期', 'latest_change_date'),
]


def build_farmland_export_workbook(view, params):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    normalized_params = {key: params.get(key) for key in params} if hasattr(params, 'get') else dict(params)
    if view == 'summary':
        sheet.title = '耕地按户汇总'
        columns = FARMLAND_HOUSEHOLD_EXPORT_COLUMNS
        data = list_farmland_households({**normalized_params, 'page': 1, 'page_size': 100000})['items']
    else:
        sheet.title = '耕地明细'
        columns = FARMLAND_DETAIL_EXPORT_COLUMNS
        data = list_farmland_records({**normalized_params, 'page': 1, 'page_size': 100000})['items']
    for col_index, (label, _) in enumerate(columns, start=1):
        sheet.cell(row=1, column=col_index, value=label)
    for row_index, item in enumerate(data, start=2):
        for col_index, (_, key) in enumerate(columns, start=1):
            sheet.cell(row=row_index, column=col_index, value=item.get(key, ''))
    return workbook


TODO_TYPE_OPTIONS = {choice[0] for choice in TodoReminder.TYPE_CHOICES}
TODO_PROGRESS_OPTIONS = {choice[0] for choice in TodoReminder.PROGRESS_CHOICES}
REMINDER_RULE_CATEGORY_OPTIONS = {choice[0] for choice in ReminderRule.CATEGORY_CHOICES}


def parse_time_value(value, default=None):
    if value in (None, ''):
        return default
    if isinstance(value, time):
        return value
    raw = serialize_scalar(value)
    if not raw:
        return default
    parts = str(raw).strip().split(':')
    if len(parts) not in (2, 3):
        raise ValueError('提醒时间格式应为 HH:MM 或 HH:MM:SS')
    try:
        hour = int(parts[0])
        minute = int(parts[1])
        second = int(parts[2]) if len(parts) == 3 else 0
        return time(hour=hour, minute=minute, second=second)
    except ValueError as exc:
        raise ValueError('提醒时间格式不正确') from exc


def ensure_aware_datetime(value):
    if value is None:
        return None
    if timezone.is_naive(value):
        return timezone.make_aware(value, timezone.get_current_timezone())
    return timezone.localtime(value)


def parse_datetime_value(value):
    if value in (None, ''):
        return None
    if isinstance(value, datetime):
        return ensure_aware_datetime(value)
    raw = str(serialize_scalar(value)).strip()
    if not raw:
        return None
    try:
        parsed = datetime.fromisoformat(raw.replace('Z', '+00:00'))
    except ValueError as exc:
        raise ValueError('提醒时间格式不正确，请使用 YYYY-MM-DDTHH:MM') from exc
    return ensure_aware_datetime(parsed)


def get_todo_status_label(item):
    if item.progress == TodoReminder.PROGRESS_COMPLETED:
        return TodoReminder.PROGRESS_COMPLETED
    return '已读' if item.is_read else '未读'


def serialize_todo_reminder(item, index=None):
    reminder_at = timezone.localtime(item.reminder_at) if item.reminder_at else None
    is_due = bool(reminder_at and reminder_at <= timezone.localtime(timezone.now()) and item.progress != TodoReminder.PROGRESS_COMPLETED)
    return {
        'id': item.id,
        'seq': index,
        'title': item.title,
        'content': item.content,
        'reminder_type': item.reminder_type,
        'progress': item.progress,
        'status': get_todo_status_label(item),
        'reminder_at': reminder_at.isoformat() if reminder_at else '',
        'reminder_display': reminder_at.strftime('%Y-%m-%d %H:%M:%S') if reminder_at else '',
        'is_read': item.is_read,
        'is_due': is_due,
        'source_type': item.source_type,
        'notes': item.notes,
        'created_at': timezone.localtime(item.created_at).strftime('%Y-%m-%d %H:%M:%S'),
        'updated_at': timezone.localtime(item.updated_at).strftime('%Y-%m-%d %H:%M:%S'),
    }


def build_todo_queryset(params):
    queryset = TodoReminder.objects.all()
    keyword = serialize_scalar(params.get('keyword'))
    if keyword:
        queryset = queryset.filter(Q(title__icontains=keyword) | Q(content__icontains=keyword))
    reminder_type = serialize_scalar(params.get('reminder_type'))
    if reminder_type and reminder_type != '全部':
        queryset = queryset.filter(reminder_type=reminder_type)
    view = serialize_scalar(params.get('view'))
    if view == '未读':
        queryset = queryset.filter(is_read=False).exclude(progress=TodoReminder.PROGRESS_COMPLETED)
    elif view in TODO_TYPE_OPTIONS:
        queryset = queryset.filter(reminder_type=view)
    status = serialize_scalar(params.get('status'))
    if status == '未读':
        queryset = queryset.filter(is_read=False).exclude(progress=TodoReminder.PROGRESS_COMPLETED)
    elif status == '已读':
        queryset = queryset.filter(is_read=True).exclude(progress=TodoReminder.PROGRESS_COMPLETED)
    elif status == TodoReminder.PROGRESS_COMPLETED:
        queryset = queryset.filter(progress=TodoReminder.PROGRESS_COMPLETED)
    progress = serialize_scalar(params.get('progress'))
    if progress and progress != '全部':
        queryset = queryset.filter(progress=progress)
    return queryset.order_by('is_read', 'reminder_at', '-created_at', '-id')


def get_todo_summary_payload(queryset=None):
    queryset = queryset if queryset is not None else TodoReminder.objects.all()
    now = timezone.localtime(timezone.now())
    total_count = queryset.count()
    unread_count = queryset.filter(is_read=False).exclude(progress=TodoReminder.PROGRESS_COMPLETED).count()
    read_count = queryset.filter(is_read=True).exclude(progress=TodoReminder.PROGRESS_COMPLETED).count()
    task_count = queryset.filter(reminder_type=TodoReminder.TYPE_TASK).count()
    not_started_count = queryset.filter(progress=TodoReminder.PROGRESS_NOT_STARTED).count()
    in_progress_count = queryset.filter(progress=TodoReminder.PROGRESS_IN_PROGRESS).count()
    completed_count = queryset.filter(progress=TodoReminder.PROGRESS_COMPLETED).count()
    due_count = queryset.filter(
        is_read=False,
        reminder_at__isnull=False,
        reminder_at__lte=now,
    ).exclude(progress=TodoReminder.PROGRESS_COMPLETED).count()
    return {
        'total_count': total_count,
        'unread_count': unread_count,
        'read_count': read_count,
        'task_count': task_count,
        'not_started_count': not_started_count,
        'in_progress_count': in_progress_count,
        'completed_count': completed_count,
        'due_count': due_count,
    }


def create_or_update_generated_todo(source_type, source_identifier, defaults):
    item = TodoReminder.objects.filter(source_identifier=source_identifier).first()
    if item is None:
        return TodoReminder.objects.create(
            source_type=source_type,
            source_identifier=source_identifier,
            **defaults,
        )
    item.title = defaults['title']
    item.content = defaults.get('content', '')
    item.reminder_type = defaults.get('reminder_type', item.reminder_type)
    item.reminder_at = defaults.get('reminder_at')
    item.notes = defaults.get('notes', '')
    item.source_type = source_type
    item.save(update_fields=['title', 'content', 'reminder_type', 'reminder_at', 'notes', 'source_type', 'updated_at'])
    return item


def cleanup_disabled_birthday_reminders():
    TodoReminder.objects.filter(
        Q(reminder_type=TodoReminder.TYPE_BIRTHDAY) | Q(source_type=TodoReminder.SOURCE_BIRTHDAY_RULE)
    ).delete()
    ReminderRule.objects.filter(category=ReminderRule.CATEGORY_BIRTHDAY).delete()


def ensure_default_reminder_rules():
    if not ReminderRule.objects.filter(category=ReminderRule.CATEGORY_PARTY_FEE).exists():
        ReminderRule.objects.create(
            category=ReminderRule.CATEGORY_PARTY_FEE,
            rule_name='党费提醒',
            reminder_time=time(9, 0, 0),
            is_month_end=True,
            is_enabled=True,
        )


def sync_birthday_reminders():
    today = timezone.localdate()
    now = timezone.localtime(timezone.now())
    current_year = today.year
    default_time = time(9, 0, 0)
    rules = ReminderRule.objects.filter(category=ReminderRule.CATEGORY_BIRTHDAY, is_enabled=True)
    residents = Resident.objects.filter(birth_date__isnull=False, status='正常')
    for rule in rules:
        if not rule.age_value:
            continue
        for resident in residents:
            birth_date = resident.birth_date
            if birth_date is None:
                continue
            try:
                birthday_this_year = birth_date.replace(year=current_year)
            except ValueError:
                birthday_this_year = birth_date.replace(year=current_year, day=28)
            target_age = birthday_this_year.year - birth_date.year
            if target_age != rule.age_value:
                continue
            remind_date = birthday_this_year - timedelta(days=rule.remind_days or 0)
            remind_at = ensure_aware_datetime(datetime.combine(remind_date, rule.reminder_time or default_time))
            if remind_at > now:
                continue
            create_or_update_generated_todo(
                TodoReminder.SOURCE_BIRTHDAY_RULE,
                f'birthday:{rule.id}:{resident.id}:{current_year}',
                {
                    'title': f'{resident.full_name}{target_age}岁生日提醒',
                    'content': f'{resident.full_name}将于{birthday_this_year.strftime("%Y-%m-%d")}满{target_age}周岁，请及时联系关怀。',
                    'reminder_type': TodoReminder.TYPE_BIRTHDAY,
                    'progress': TodoReminder.PROGRESS_NOT_STARTED,
                    'reminder_at': remind_at,
                    'notes': resident.phone or '',
                },
            )


def sync_party_fee_reminders():
    today = timezone.localdate()
    now = timezone.localtime(timezone.now())
    default_time = time(9, 0, 0)
    rules = ReminderRule.objects.filter(category=ReminderRule.CATEGORY_PARTY_FEE, is_enabled=True)
    last_day = monthrange(today.year, today.month)[1]
    records = PartyFeeRecord.objects.filter(fee_year=today.year, fee_month=today.month)
    pending_count = records.filter(payment_status=PartyFeeRecord.PAYMENT_STATUS_PENDING).count()
    total_count = records.count()
    for rule in rules:
        remind_day = last_day if rule.is_month_end or not rule.reminder_day else min(rule.reminder_day, last_day)
        remind_at = ensure_aware_datetime(datetime.combine(today.replace(day=remind_day), rule.reminder_time or default_time))
        if remind_at > now:
            continue
        if total_count == 0:
            title = f'{today.year}年{today.month}月党费清单提醒'
            content = f'{today.year}年{today.month}月党费清单尚未生成，请先生成当月党费记录。'
        elif pending_count > 0:
            title = f'{today.year}年{today.month}月党费缴纳提醒'
            content = f'当前仍有 {pending_count} 条党费待缴纳，请及时处理。'
        else:
            title = f'{today.year}年{today.month}月党费缴纳提醒'
            content = f'{today.year}年{today.month}月党费已全部缴纳。'
        create_or_update_generated_todo(
            TodoReminder.SOURCE_PARTY_FEE_RULE,
            f'party-fee:{rule.id}:{today.year}:{today.month}',
            {
                'title': title,
                'content': content,
                'reminder_type': TodoReminder.TYPE_SYSTEM,
                'progress': TodoReminder.PROGRESS_NOT_STARTED,
                'reminder_at': remind_at,
                'notes': rule.rule_name,
            },
        )


def sync_generated_todo_reminders():
    cleanup_disabled_birthday_reminders()
    ensure_default_reminder_rules()
    sync_party_fee_reminders()


@transaction.atomic
def list_todo_reminders(params):
    sync_generated_todo_reminders()
    page_number = int(params.get('page', 1) or 1)
    page_size = int(params.get('page_size', 10) or 10)
    queryset = build_todo_queryset(params)
    stats = get_todo_summary_payload(queryset)
    paginator, page = paginate_queryset(queryset, page_number, page_size)
    start_index = (page.number - 1) * page_size + 1
    items = [serialize_todo_reminder(item, start_index + offset) for offset, item in enumerate(page.object_list)]
    return {
        'items': items,
        'pagination': {
            'page': page.number,
            'page_size': page_size,
            'total': paginator.count,
            'total_pages': paginator.num_pages,
        },
        'stats': stats,
    }


def validate_todo_payload(data):
    title = serialize_scalar(data.get('title'))
    if not title:
        raise ValueError('标题不能为空')
    reminder_type = serialize_scalar(data.get('reminder_type')) or TodoReminder.TYPE_TASK
    if reminder_type not in TODO_TYPE_OPTIONS:
        raise ValueError('提醒类型不正确')
    progress = serialize_scalar(data.get('progress')) or TodoReminder.PROGRESS_NOT_STARTED
    if progress not in TODO_PROGRESS_OPTIONS:
        raise ValueError('进度状态不正确')
    return {
        'title': title,
        'content': serialize_scalar(data.get('content')),
        'reminder_type': reminder_type,
        'progress': progress,
        'reminder_at': parse_datetime_value(data.get('reminder_at')),
        'notes': serialize_scalar(data.get('notes')),
        'is_read': bool(data.get('is_read')),
    }


@transaction.atomic
def create_todo_reminder(data):
    payload = validate_todo_payload(data)
    item = TodoReminder.objects.create(
        title=payload['title'],
        content=payload['content'],
        reminder_type=payload['reminder_type'],
        progress=payload['progress'],
        reminder_at=payload['reminder_at'],
        notes=payload['notes'],
        is_read=payload['is_read'],
        read_at=timezone.now() if payload['is_read'] else None,
        source_type=TodoReminder.SOURCE_MANUAL,
    )
    return serialize_todo_reminder(item)


def get_todo_reminder_detail(record_id):
    item = TodoReminder.objects.get(id=record_id)
    return serialize_todo_reminder(item)


@transaction.atomic
def update_todo_reminder(record_id, data):
    item = TodoReminder.objects.get(id=record_id)
    payload = validate_todo_payload(data)
    item.title = payload['title']
    item.content = payload['content']
    item.reminder_type = payload['reminder_type']
    item.progress = payload['progress']
    item.reminder_at = payload['reminder_at']
    item.notes = payload['notes']
    if payload['is_read']:
        item.is_read = True
        item.read_at = item.read_at or timezone.now()
    else:
        item.is_read = False
        item.read_at = None
    item.save()
    return serialize_todo_reminder(item)


@transaction.atomic
def delete_todo_reminder(record_id):
    item = TodoReminder.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}


@transaction.atomic
def bulk_delete_todo_reminders(ids):
    deleted, _ = TodoReminder.objects.filter(id__in=list(ids or [])).delete()
    return {'message': f'成功删除 {deleted} 条记录'}


@transaction.atomic
def bulk_update_todo_read_status(ids, is_read):
    queryset = TodoReminder.objects.filter(id__in=list(ids or []))
    if is_read:
        count = queryset.update(is_read=True, read_at=timezone.now(), updated_at=timezone.now())
        return {'message': f'成功将 {count} 条提醒设为已读'}
    count = queryset.update(is_read=False, read_at=None, updated_at=timezone.now())
    return {'message': f'成功将 {count} 条提醒设为未读'}


@transaction.atomic
def mark_all_todo_read(params=None):
    queryset = build_todo_queryset(params or {})
    count = queryset.filter(is_read=False).update(is_read=True, read_at=timezone.now(), updated_at=timezone.now())
    return {'message': f'成功将 {count} 条提醒设为已读'}


def get_todo_summary():
    sync_generated_todo_reminders()
    summary = get_todo_summary_payload(TodoReminder.objects.all())
    summary['latest_items'] = [
        serialize_todo_reminder(item)
        for item in TodoReminder.objects.filter(is_read=False).exclude(progress=TodoReminder.PROGRESS_COMPLETED).order_by('reminder_at', '-created_at')[:5]
    ]
    return summary


def serialize_reminder_rule(item):
    reminder_time_display = item.reminder_time.strftime('%H:%M:%S') if item.reminder_time else ''
    if item.category == ReminderRule.CATEGORY_BIRTHDAY:
        reminder_date_label = ''
        age_condition_label = f'{item.age_value}岁' if item.age_value else ''
    else:
        reminder_date_label = '每月月底' if item.is_month_end else f'每月{item.reminder_day}日'
        age_condition_label = ''
    return {
        'id': item.id,
        'category': item.category,
        'rule_name': item.rule_name,
        'age_value': item.age_value,
        'age_condition_label': age_condition_label,
        'remind_days': item.remind_days,
        'reminder_time': reminder_time_display,
        'reminder_day': item.reminder_day,
        'reminder_date_label': reminder_date_label,
        'is_month_end': item.is_month_end,
        'is_enabled': item.is_enabled,
        'notes': item.notes,
    }


def list_reminder_rules():
    cleanup_disabled_birthday_reminders()
    ensure_default_reminder_rules()
    queryset = ReminderRule.objects.all()
    return {
        'birthday_rules': [],
        'party_fee_rules': [
            serialize_reminder_rule(item)
            for item in queryset.filter(category=ReminderRule.CATEGORY_PARTY_FEE)
        ],
    }


def validate_reminder_rule_payload(data, current=None):
    category = serialize_scalar(data.get('category')) or (current.category if current else ReminderRule.CATEGORY_PARTY_FEE)
    if category not in REMINDER_RULE_CATEGORY_OPTIONS:
        raise ValueError('提醒规则分类不正确')
    if category == ReminderRule.CATEGORY_BIRTHDAY:
        raise ValueError('生日提醒已取消，不再支持配置')
    rule_name = serialize_scalar(data.get('rule_name')) or '党费提醒'
    reminder_time = parse_time_value(data.get('reminder_time'), default=time(9, 0, 0))
    is_enabled = bool(data.get('is_enabled', True if current is None else current.is_enabled))
    notes = serialize_scalar(data.get('notes'))
    reminder_day = parse_int_value(data.get('reminder_day'), 0)
    is_month_end = bool(data.get('is_month_end')) or reminder_day <= 0
    if not is_month_end and (reminder_day < 1 or reminder_day > 31):
        raise ValueError('党费提醒日期必须在 1 到 31 之间')
    return {
        'category': category,
        'rule_name': rule_name,
        'age_value': None,
        'remind_days': 0,
        'reminder_time': reminder_time,
        'reminder_day': None if is_month_end else reminder_day,
        'is_month_end': is_month_end,
        'is_enabled': is_enabled,
        'notes': notes,
    }


@transaction.atomic
def create_reminder_rule(data):
    payload = validate_reminder_rule_payload(data)
    item = ReminderRule.objects.create(**payload)
    return serialize_reminder_rule(item)


def get_reminder_rule_detail(record_id):
    item = ReminderRule.objects.get(id=record_id)
    return serialize_reminder_rule(item)


@transaction.atomic
def update_reminder_rule(record_id, data):
    item = ReminderRule.objects.get(id=record_id)
    payload = validate_reminder_rule_payload(data, current=item)
    for field, value in payload.items():
        setattr(item, field, value)
    item.save()
    return serialize_reminder_rule(item)


@transaction.atomic
def delete_reminder_rule(record_id):
    item = ReminderRule.objects.get(id=record_id)
    item.delete()
    return {'message': '删除成功'}
